"""Build the customer-facing style-divergence Word document.

A send-ready deliverable: how the NIEHS-supplied .dotx templates compare to the
style library the automated generator uses (extracted from the NIEHS-10 reference
report), why the divergences exist (manual-preparation forensics), and what we
recommend.  Facts are the verified figures from the concordance sweep + the
style-definition forensics; this script only lays them out.

Not shipped -- a one-off deliverable builder.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches

OUT = Path("output") / "NIEHS-Style-Template-Divergences.docx"

NAVY = RGBColor(0x1F, 0x38, 0x63)
GREY = RGBColor(0x59, 0x59, 0x59)
GREEN = RGBColor(0x2E, 0x6B, 0x2E)
AMBER = RGBColor(0x9A, 0x6A, 0x00)


def _shade(cell, hex_fill):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _keep_with_next(p):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    kwn = OxmlElement("w:keepNext")
    pPr.append(kwn)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    _keep_with_next(p)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    _keep_with_next(p)
    return p


def body(doc, text, size=10.5, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(after)
    return p


def bullet(doc, text, size=10.5, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        rb = p.add_run(bold_lead)
        rb.bold = True
        rb.font.size = Pt(size)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def footnote(doc, segments):
    """A small-italic note under a table. `segments` is a list of (text, mono?)
    tuples; mono renders the marker/class names in a monospaced face."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    for text, mono in segments:
        r = p.add_run(text)
        r.font.size = Pt(8.5)
        r.font.color.rgb = GREY
        if mono:
            r.font.name = "Consolas"
        else:
            r.italic = True
    return p


def callout(doc, lead, text, color=AMBER):
    p = doc.add_paragraph()
    rb = p.add_run(lead)
    rb.bold = True
    rb.font.size = Pt(10)
    rb.font.color.rgb = color
    r = p.add_run(text)
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0].add_run(htext)
        rp.bold = True
        rp.font.size = Pt(9.5)
        rp.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "1F3863")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rp = cells[i].paragraphs[0].add_run(str(val))
            rp.font.size = Pt(9.5)
        if ri % 2 == 1:
            for c in cells:
                _shade(c, "F2F5FA")
    if widths:
        for col, w in zip(t.columns, widths):
            for c in col.cells:
                c.width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ================= COVER / TITLE BLOCK =================
    tp = doc.add_paragraph()
    tr = tp.add_run("Report Style Templates: Divergence Analysis")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = NAVY
    tp.paragraph_format.space_after = Pt(2)

    sp = doc.add_paragraph()
    sr = sp.add_run(
        "How the supplied Word style templates compare to the styling of the "
        "NIEHS-10 reference report, and what it means for automated generation"
    )
    sr.font.size = Pt(12)
    sr.font.color.rgb = GREY
    sr.italic = True
    sp.paragraph_format.space_after = Pt(10)

    meta = doc.add_paragraph()
    mr = meta.add_run("Prepared for:  NIEHS  ·  Subject:  Automated report styling  ·  Status:  For review and discussion")
    mr.font.size = Pt(9)
    mr.font.color.rgb = GREY

    # thin rule
    rule = doc.add_paragraph()
    rr = rule.add_run("_" * 92)
    rr.font.color.rgb = RGBColor(0xC8, 0xD2, 0xE0)
    rule.paragraph_format.space_after = Pt(8)

    # ================= EXECUTIVE SUMMARY =================
    h1(doc, "Executive summary")
    p = doc.add_paragraph()
    r = p.add_run(
        "We compared the two Word style templates NIEHS provided against the "
        "styling actually present in the NIEHS-10 reference report. The two are "
        "highly consistent in appearance, but they are not identical in "
        "composition — and understanding why is useful for the reporting program.")
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)

    bullet(doc, "the visual definition of every shared style is identical. The "
                "NIEHS look — fonts, sizes, weights, spacing — is consistent "
                "across both sources, in all 380 styles they share.",
           bold_lead="Where they agree:  ")
    bullet(doc, "the reference report contains 6 styles the supplied templates do "
                "not, and the templates contain 4 styles the reference report does "
                "not — 10 differences in total.",
           bold_lead="Where they differ:  ")
    bullet(doc, "the differences are almost entirely the fingerprint of manual "
                "document preparation. Five of the six “extra” styles in the "
                "reference report are residue left behind when web content was "
                "copied and pasted into Word during authoring — not intentional "
                "house style.",
           bold_lead="Why they differ:  ")
    bullet(doc, "the automated generator does not use the supplied templates. It "
                "uses a style library taken directly from the reference report, so "
                "its output matches that report exactly (386 of 386 styles). This "
                "is a deliberate fidelity choice.",
           bold_lead="What the generator does:  ")

    p = doc.add_paragraph()
    r = p.add_run("Bottom line:  ")
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = GREEN
    r2 = p.add_run(
        "there is no styling defect in either artifact. The divergences are the "
        "natural byproduct of a human editing process, and they are exactly the "
        "kind of variation that automated generation eliminates. The items in "
        "Section 6 are decisions to confirm with NIEHS, not problems to fix.")
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)

    # ================= 1. BACKGROUND =================
    h1(doc, "1.  Background")
    body(doc,
         "A Microsoft Word document separates content from styling. Each paragraph "
         "is tagged with a style name (for example, “3-03_Head2”); the definition "
         "of what that style looks like lives in the document's style library. A "
         "“.dotx” template is a carrier for such a style library. This report "
         "concerns the style libraries — not the words of any report.")
    body(doc, "Three artifacts are relevant:")
    table(doc,
          ["Artifact", "Origin", "Role"],
          [["Reference report", "NIEHS-10 PFHxSAm final .docx, delivered by NIEHS",
            "Ground truth for the target appearance"],
           ["Generator style library", "Extracted verbatim from the reference report",
            "Bundled into every generated report"],
           ["Supplied templates (×2)", "“…-bordered.dotx” and “…-borderless.dotx”",
            "House-style templates provided separately"]],
          widths=[1.7, 3.0, 1.9])
    body(doc,
         "The bordered and borderless templates carry an identical set of style "
         "names (384 each) and differ only in decorative page borders; they are "
         "treated as one for this analysis.", size=9.5)

    # ================= 2. WHAT WE COMPARED =================
    h1(doc, "2.  Summary of findings")
    table(doc,
          ["Comparison", "Shared styles", "Only in reference", "Only in template",
           "Definition mismatches"],
          [["Generator library  vs.  reference report", "386", "0", "0", "0"],
           ["Supplied templates  vs.  reference report", "380", "6", "4", "0"]],
          widths=[3.0, 1.1, 1.3, 1.2, 1.4])
    body(doc,
         "Read the top row as: the generator reproduces the reference report's "
         "styling exactly. Read the bottom row as: the supplied templates differ "
         "from the reference report by 10 styles, but never disagree about what a "
         "shared style should look like.")

    # ================= 3. THE DIVERGENCES =================
    h1(doc, "3.  The divergences in detail")

    h2(doc, "3.1  In the reference report, but not in the supplied templates (6)")
    body(doc,
         "These styles are present in the delivered report but absent from the "
         "house templates:", size=10, after=4)
    table(doc,
          ["Style name", "Nature", "Notes"],
          [["No Spacing", "Word built-in", "Standard Word style; harmless."],
           ["normaltextrun", "Web-paste residue", "A Microsoft web-editor CSS class."],
           ["dx-doi", "Web-paste residue", "DOI-link styling from a publisher web page."],
           ["dictionarylistcontent-951", "Web-paste residue", "Reference-site class; carries the “Segoe UI” web font."],
           ["pcon", "Web-paste residue", "Browser-origin class, no formatting."],
           ["sbuls", "Web-paste residue", "Browser-origin class, no formatting."]],
          widths=[2.2, 1.7, 2.7])
    footnote(doc, [
        ("How origin was determined:  ", False),
        ("the five “web-paste residue” styles are custom styles whose names are "
         "CSS classes from web pages (", False),
        ("normaltextrun", True),
        (" is emitted by Microsoft's web editors; ", False),
        ("dx-doi", True),
        (" and ", False),
        ("dictionarylistcontent", True),
        (" come from publisher / reference sites), and the three paragraph styles "
         "among them carry the marker ", False),
        ("w:beforeAutospacing=\"1\"", True),
        (", which Word writes only when importing an HTML paragraph — a reliable "
         "signature of pasted web content. “No Spacing”, by contrast, is a "
         "standard Word built-in (not a custom style and not web-derived); it is "
         "listed here only because it, too, is unused and is not part of "
         "intentional NIEHS house style.", False),
    ])

    h2(doc, "3.2  In the supplied templates, but not in the reference report (4)")
    table(doc,
          ["Style name", "Nature", "Notes"],
          [["Title", "Word built-in", "The reference uses the custom “1-03_Report_Title” instead — see note."],
           ["Title Char", "Companion character style", "Paired with Title."],
           ["TOC Heading", "Word built-in", "Table-of-contents heading."],
           ["nlp", "Template-specific", "Purpose unknown; to be confirmed with NIEHS."]],
          widths=[2.2, 1.8, 2.6])
    callout(doc, "On “Title”:  ",
            "the templates define Word's built-in Title style, but the reference "
            "report does not use it — its title is set with the custom NIEHS style "
            "“1-03_Report_Title”, which is present and identical in both sources. "
            "This only matters if a future document relies on the built-in Title; "
            "the generator follows the reference report's convention.")

    h2(doc, "3.3  Where the two sources agree completely")
    p = doc.add_paragraph()
    r = p.add_run("For all 380 shared styles, the visual definition is identical.")
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = GREEN
    p.paragraph_format.space_after = Pt(4)
    body(doc,
         "Checked across font, size, color, weight, emphasis, justification, and "
         "paragraph spacing. Representative core styles:", size=10, after=4)
    table(doc,
          ["Style", "Definition (identical in both)"],
          [["1-03_Report_Title", "20 pt, bold"],
           ["3-02a_Head1_NoNumber", "17 pt, bold"],
           ["3-03_Head2", "15 pt, bold"],
           ["3-04_Head3", "13 pt, bold"],
           ["0-03_Paragraph", "body paragraph, regular"],
           ["Normal", "base default"]],
          widths=[2.6, 4.0])
    body(doc,
         "The theme (color and font scheme) and page geometry — US Letter with "
         "one-inch margins — are also identical to the reference report. The "
         "disagreement is only ever about which styles exist, never about how a "
         "shared style looks.", size=10)

    # ================= 4. WHY MANUAL PREPARATION CAUSES DIVERGENCE =================
    h1(doc, "4.  Why manual document preparation causes these divergences")
    body(doc,
         "The divergences are not random. Their pattern is the signature of a "
         "human editing process, and the style definitions themselves carry the "
         "evidence.")

    h2(doc, "4.1  The evidence: copy-paste residue")
    body(doc,
         "Five of the six “extra” styles in the reference report — normaltextrun, "
         "dx-doi, dictionarylistcontent-951, pcon, sbuls — are not report styles "
         "at all. They are CSS class names from web pages. “normaltextrun” is "
         "emitted by Microsoft's own web editors; “dx-doi” and "
         "“dictionarylistcontent” come from academic-publisher and reference "
         "websites; the last carries “Segoe UI”, the Windows web-interface font "
         "that no one would choose for a printed report. (The sixth style, “No "
         "Spacing”, is not web residue but an incidental Word built-in; it is "
         "simply unused.)")
    body(doc,
         "The origin is not merely inferred from the names. Three of the pasted "
         "styles carry an internal marker — “beforeAutospacing” — that Microsoft "
         "Word writes only when it converts an HTML paragraph into a Word "
         "paragraph. Native Word authoring never produces it. Its presence is a "
         "reliable signature that the text arrived by pasting from a web page.")
    body(doc,
         "When an author copies text from a web page and pastes it into Word with "
         "formatting kept, Word imports the source page's CSS classes as Word "
         "styles to preserve the look. Crucially, those style definitions remain in "
         "the file even after the pasted text is deleted or restyled — Word does "
         "not remove unused styles. The five web styles are therefore fossils of "
         "authoring sessions in which source material (abstracts, DOI links, "
         "reference lists) was pasted into the draft.")

    h2(doc, "4.2  The general pattern")
    bullet(doc, "each “gather sources and paste” action can inject foreign styles "
                "that persist invisibly. A hand-assembled report accumulates them "
                "over its life; a generated report never does.",
           bold_lead="Paste contamination:  ")
    bullet(doc, "under deadline, an author eventually meets a case the house styles "
                "don't quite cover and creates a one-off style or applies manual "
                "formatting. Across authors, reviewers, and revision rounds, the "
                "document's style set drifts from the template it began with.",
           bold_lead="Style drift:  ")
    bullet(doc, "the author opens a clean template, but from the first paste and "
                "first ad-hoc tweak the document's style library forks from the "
                "template's. The delivered file reflects its entire editing "
                "history, not the pristine template.",
           bold_lead="Template ≠ delivered document:  ")

    h2(doc, "4.3  Reading the direction of a divergence")
    body(doc,
         "The direction of each difference points to its cause, which makes it a "
         "useful diagnostic:", size=10, after=4)
    table(doc,
          ["Difference", "What it indicates"],
          [["Present in the report, absent from the template",
            "Something the editing process ADDED — paste residue or ad-hoc styles. "
            "Accumulated bloat from manual preparation."],
           ["Present in the template, absent from the report",
            "Template capacity this particular author did not use, or a template "
            "revised after the report was authored. Not a manual-prep artifact."]],
          widths=[3.0, 3.6])

    h2(doc, "4.4  Why this favors automated generation")
    body(doc,
         "Manual preparation is entropic: each report is a fresh accumulation of "
         "paste residue, drift, and one-off decisions, so no two hand-made reports "
         "are stylistically identical even when they start from the same template. "
         "Automated generation is deterministic: it emits only the styles it is "
         "given, from a fixed library, every time. It cannot paste from a browser, "
         "cannot tire and improvise a heading, and cannot leave an orphaned style "
         "behind. The absence of these divergences in generated output is a "
         "concrete quality benefit, not an accident.")

    # ================= 5. WHAT THE GENERATOR USES =================
    h1(doc, "5.  What the automated generator uses, and why")
    body(doc,
         "The generator does not attach or depend on the supplied templates. For "
         "each report it writes the correct NIEHS style name onto every element "
         "and bundles a style library taken directly from the reference report. "
         "Because that library is the reference report's own, generated output "
         "matches the reference exactly (386 of 386 styles).")
    callout(doc, "Note:  ",
            "adopting the supplied templates as the style source would actually "
            "reduce fidelity to the reference report — it would introduce the 4 "
            "template-only styles and drop the reference's own — so we retain the "
            "reference-derived library. The single best future input would be the "
            "exact template the reference report was authored from.",
            color=NAVY)

    # ================= 6. FOR DISCUSSION =================
    h1(doc, "6.  Recommended decisions for NIEHS")
    bullet(doc, "Confirm that the 5 web-paste residue styles (normaltextrun, dx-doi, "
                "dictionarylistcontent-951, pcon, sbuls) are not intentional house "
                "style and may be dropped from the generator's library for cleanliness.",
           bold_lead="1.  ")
    bullet(doc, "Decide the report-title convention: the custom "
                "“1-03_Report_Title” (as the reference report uses) versus the "
                "built-in “Title” (as the supplied templates define). The generator "
                "currently follows the reference report.",
           bold_lead="2.  ")
    bullet(doc, "Clarify the purpose of the template-only “nlp” style.",
           bold_lead="3.  ")
    bullet(doc, "If available, provide the exact .dotx from which the reference "
                "report was authored. This would be the cleanest single source of "
                "truth — free of paste residue and with the intended title "
                "convention — and could replace the report-extracted library.",
           bold_lead="4.  ")

    # ================= APPENDIX =================
    h1(doc, "Appendix.  Method")
    body(doc,
         "Findings were produced by programmatically extracting the style library "
         "(styles.xml) from each file and comparing them by style name and by "
         "definition. Definition comparison covered font, size, color, bold, "
         "italic, justification, and paragraph spacing for every shared style. "
         "Style origin (built-in, custom, or web-paste) was determined from each "
         "style's definition attributes. Theme, page geometry, numbering, font "
         "table, and document settings were compared as supplementary parts. All "
         "counts in this document are exact, not estimates.", size=9.5)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
