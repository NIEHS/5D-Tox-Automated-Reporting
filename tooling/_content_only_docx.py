"""Test: content/style separation via "attach template after".

Takes the BODY CONTENT of the reference dotx (discarding its style DEFINITIONS)
and re-wraps it in a fresh python-docx Document -- i.e. MS Word DEFAULT styles,
NO NIEHS definitions.  Every paragraph still REFERENCES a built-in style name
(Title / Heading 1..3 / Normal); those names exist in the docx but with Word's
plain default look.

Intended manual workflow: open the resulting docx in OnlyOffice/Word, then attach
`niehs-10-base.dotx` with "automatically update document styles" -- Word matches
each paragraph's style NAME to the template's definition and restyles the content
to the NIEHS look, without the content file ever carrying those definitions.

Not shipped -- a separation test harness.
"""
from pathlib import Path
import zipfile
import re

from docx import Document

REF_DOTX = Path("assets/templates/NIEHS-report-style-borderless.dotx")
OUT = Path("output") / "content-only-worddefault.docx"

# python-docx built-in style NAME per reference styleId (styleId -> add spec).
# level is used for headings; None => plain paragraph with that style name.
_STYLEID_TO_BUILTIN = {
    "Title": ("Title", None),
    "Heading1": ("Heading 1", 1),
    "Heading2": ("Heading 2", 2),
    "Heading3": ("Heading 3", 3),
}


def extract_content(path: Path):
    """Return [(styleId_or_None, text)] for each non-empty body paragraph."""
    xml = zipfile.ZipFile(path).read("word/document.xml").decode("utf-8")
    out = []
    for p in re.findall(r"<w:p\b.*?</w:p>", xml, re.S):
        ps = re.search(r'<w:pStyle w:val="([^"]+)"', p)
        text = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", p)).strip()
        if text:
            out.append((ps.group(1) if ps else None, text))
    return out


def build() -> None:
    if not REF_DOTX.exists():
        raise SystemExit(f"missing reference dotx: {REF_DOTX}")

    paras = extract_content(REF_DOTX)

    # Fresh Document() == MS Word defaults; carries NO NIEHS style definitions.
    doc = Document()

    for style_id, text in paras:
        if style_id in _STYLEID_TO_BUILTIN:
            name, level = _STYLEID_TO_BUILTIN[style_id]
            if level is not None:
                doc.add_heading(text, level=level)
            else:
                doc.add_paragraph(text, style=name)
        else:
            # No pStyle in the source => Normal (Word default body).
            doc.add_paragraph(text)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")
    print(f"  {len(paras)} content paragraphs, referencing built-in style names:")
    for style_id, text in paras:
        label = style_id or "Normal"
        print(f"    [{label}] {text[:60]}")


if __name__ == "__main__":
    build()
