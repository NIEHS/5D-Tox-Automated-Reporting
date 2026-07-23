r"""
docx_style_extract.py — bootstrap the styling vocabulary FROM a Word template.

The inverse of docx_generator's style application: open a `.docx` authored in
Word/LibreOffice and read its base styles + page geometry into the `styles`
mapping the rest of the pipeline already consumes (the {defaults, types,
instances, document} shape `document_config._parse_styles_yaml` validates and
`report_data._resolve_layout_config` feeds to all three renderers).

This is how a customer designs the report's look in Word — WYSIWYG — and has it
drive the LaTeX and HTML surfaces too, at the shared-vocabulary fidelity
(layout_style.LAYOUT_KEY_SCHEMA + DOCUMENT_KEY_SCHEMA — the ~20 typographic +
document-geometry keys).

SCOPE — buckets 1 + 2 only (see the design discussion):
  1. per-block typography — the Word STYLES Normal / Heading 1-3 / Header →
     the `defaults` + `types` layers.
  2. document geometry — the first section's page size, margins, header
     distance, and the base body/header fonts → the `document` layer.
Bucket 3 (table column widths, ToC dot-leaders, list numbering — Word's
COMPONENT geometry) is deliberately NOT extracted; it isn't part of the shared
vocabulary and stays per-surface. `extract_styles` logs what it skipped so the
omission is visible, never silent.

Font names are LITERAL on every surface (the decision for this feature): a
"Times New Roman" read here is written verbatim as the `font` key and used as-is
by docx, LaTeX (fontspec), and HTML — not mapped back to serif/sans/mono.

Read-only w.r.t. fonts: this reads a template's style *names* (family strings),
never glyph data, so it carries no font-licensing burden.

CLI:
    python -m docx_style_extract template.docx            # YAML to stdout
    python -m docx_style_extract template.docx -o out.yaml
"""

from __future__ import annotations

import io
import re
import sys
import zipfile

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Emu

# python-docx opens .docx (document.main) but rejects .dotx (template.main).
# A template is the same OPC package with one content-type string changed, so we
# normalize it in memory to read a Word template the customer authored as .dotx.
_TEMPLATE_CT = b"wordprocessingml.template.main+xml"
_DOCUMENT_CT = b"wordprocessingml.document.main+xml"


def _open_word(path: str):
    """Open a .docx OR .dotx into a python-docx Document (template CT normalized)."""
    with zipfile.ZipFile(path) as zin:
        if _TEMPLATE_CT not in zin.read("[Content_Types].xml"):
            return Document(path)  # plain .docx — open directly
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.replace(_TEMPLATE_CT, _DOCUMENT_CT)
                zout.writestr(item, data)
    buf.seek(0)
    return Document(buf)

# The Word STYLE name → the styles-config `types` node_type it maps to.  Normal
# is the document default (→ `defaults`), Header feeds the `document` layer's
# header font; the three Heading styles map to the heading-bearing catalog types
# that share a level.  We key on the built-in style names docx_generator emits.
_HEADING_STYLE_TO_TYPES: dict[str, list[str]] = {
    # Heading 1 → the level-1 body sections; Heading 2/3 → the nested groups.
    # These node_type names must exist in the component catalog (validated
    # downstream); we map to the ones docx_generator styles via _HEADING_STYLE_BY_LEVEL.
    "Heading 1": ["front-matter", "narrative", "heading-only", "appendix",
                  "bmd-summary", "genomics-section"],
    "Heading 2": ["narrative+tables"],
    "Heading 3": [],  # level-3 nodes inherit Heading 3 in Word; no distinct type today
}

# The NTP template's title-page style family → the title_page ROLE each maps to
# (layout_style.TITLE_PAGE_ROLES).  The style NAMES are the ICF/NTP `1-NN`
# convention (Word displays them with a space + underscores).  Only the roles
# that correspond to content OUR title-page node emits are mapped; the rest of
# the `1-NN` family (foreword/abstract/authors/peer-review) belongs to separate
# front-matter nodes and is intentionally not read here.
_TITLE_PAGE_STYLE_TO_ROLE: dict[str, str] = {
    "1-03_Report_Title": "report_title",
    "1-02_Report_Type": "report_type",
    "1-04_Report_Subtitle": "report_subtitle",
    "1-05_Publication_Date": "publication_date",
    "1-05a_Report-Number": "report_number",
    "1-05b_DOI": "doi",
    "1-05c_ISSN": "issn",
    "1-05d_NIH-Number": "nih_number",
    "1-01_Publisher_Name": "publisher_name",
    "1-06_Publication_Office": "publication_office",
    "1-07_Publication_Division": "publication_division",
    "1-08_Publication_Institute": "publication_institute",
    "1-09_Publication_Department": "publication_department",
    "1-26_Logo_Graphic": "logo_graphic",
}


# ---------------------------------------------------------------------------
# Length / color helpers — python-docx returns EMU Length objects and RGBColor
# ---------------------------------------------------------------------------

_EMU_PER_PT = 12700  # 1pt = 12700 EMU


def _emu_to_pt_str(length) -> "str | None":
    """A python-docx Length (EMU) → a 'Npt' string the vocabulary accepts, or None."""
    if length is None:
        return None
    pt = int(length) / _EMU_PER_PT
    # Trim trailing-zero noise: 12.0 → "12pt", 13.8 → "13.8pt".
    s = f"{pt:.2f}".rstrip("0").rstrip(".")
    return f"{s}pt"


def _emu_to_in_str(length) -> "str | None":
    """A python-docx Length (EMU) → an 'Nin' string (used for page/margin sizes)."""
    if length is None:
        return None
    inches = int(length) / 914400
    s = f"{inches:.3f}".rstrip("0").rstrip(".")
    return f"{s}in"


def _color_to_hex(font) -> "str | None":
    """A style font's color → '#rrggbb', or None when unset/auto/theme."""
    color = font.color
    if color is None or color.type is None:
        return None
    rgb = color.rgb
    if rgb is None:
        return None
    return f"#{str(rgb).lower()}"


def _rpr_letter_spacing_twips(font) -> "int | None":
    """Character spacing from a font's rPr <w:spacing w:val> (twips), or None.

    This is the rPr (run-property) w:spacing — inter-character tracking — NOT the
    pPr w:spacing (before/after paragraph space).  python-docx's Font has no
    accessor for it, so we read the element directly.  Returns None when the style
    sets no character spacing (so the key stays absent); 0 is a real value."""
    rpr = getattr(font, "element", None)
    rpr = getattr(rpr, "rPr", None)
    if rpr is None:
        return None
    sp = rpr.find(qn("w:spacing"))
    if sp is None:
        return None
    val = sp.get(qn("w:val"))
    if val is None:
        return None
    try:
        return int(val)
    except ValueError:
        return None


# ---------------------------------------------------------------------------
# Per-style extraction → one resolved-style dict (the vocabulary's flat keys)
# ---------------------------------------------------------------------------

def _extract_style_props(style) -> dict:
    """Read one Word paragraph style into a vocabulary style dict (only the keys
    the style actually sets — absent props stay absent so layers stay sparse)."""
    out: dict = {}
    font = style.font
    if font.name:
        out["font"] = font.name           # literal family name (font wins)
    size = _emu_to_pt_str(font.size)
    if size:
        out["font_size"] = size
    if font.bold is True:
        out["weight"] = "bold"
    elif font.bold is False:
        out["weight"] = "normal"
    if font.italic is True:
        out["style"] = "italic"
    elif font.italic is False:
        out["style"] = "normal"
    # w:caps (all-caps display) → text_transform.  Only uppercase is in the
    # vocabulary; w:smallCaps has no faithful cross-surface mapping so it is not
    # read.  font.all_caps is None when unset → key stays absent (sparse layer).
    if font.all_caps is True:
        out["text_transform"] = "uppercase"
    elif font.all_caps is False:
        out["text_transform"] = "none"
    # letter_spacing ← rPr <w:spacing w:val> (twips → pt).  CRITICAL: this is the
    # CHARACTER-spacing element inside rPr, NOT paragraph pPr <w:spacing
    # w:before/after> (that's space_before/after, read from paragraph_format
    # below).  We read it off the font (rPr) element so the two can't be confused
    # (the ADR-0009 trap).  w:val is signed twips; 0 is a real value (no spacing)
    # so we test for the attribute's presence, not truthiness.
    ls_twips = _rpr_letter_spacing_twips(font)
    if ls_twips is not None:
        out["letter_spacing"] = _emu_to_pt_str(int(ls_twips) / 20 * _EMU_PER_PT)
    color = _color_to_hex(font)
    if color:
        out["color"] = color

    pf = style.paragraph_format
    if pf.alignment is not None:
        out["align"] = {0: "left", 1: "center", 2: "right", 3: "justify"}.get(
            int(pf.alignment)
        )
        if out["align"] is None:
            del out["align"]
    # line_spacing is a float multiple only when the RULE is MULTIPLE/SINGLE/etc.;
    # an EMU value (exact spacing) can't be a unitless multiplier, so skip it.
    if isinstance(pf.line_spacing, float):
        out["line_height"] = pf.line_spacing
    sb = _emu_to_pt_str(pf.space_before)
    if sb:
        out["space_before"] = sb
    sa = _emu_to_pt_str(pf.space_after)
    if sa:
        out["space_after"] = sa
    indent = _emu_to_pt_str(pf.first_line_indent)
    if indent:
        out["first_line_indent"] = indent
    if pf.page_break_before is True:
        out["break_before"] = "page"
    # No symmetric break_after: OOXML has a pageBreakBefore paragraph property but
    # NO page-break-after.  An after-break is a page-break RUN in body content, not
    # a style property, so it is structurally unreadable from a paragraph style —
    # break_after stays emit-only (see docx_generator._layout_to_docx).
    if pf.keep_together is True:
        out["keep_together"] = True
    return out


def _resolved_style_props(style) -> dict:
    """Style props with the ``basedOn`` (parent) chain resolved.

    python-docx does NOT resolve inheritance: `style.font.name` on a style that
    inherits its font from a parent (via ``w:basedOn``) returns None.  The NTP
    title-page styles rely on this — e.g. `1-03_Report_Title` gets its Arial from
    the parent `Base_Heading` — so reading a child style alone loses the font.
    We walk `style.base_style` to the root, then merge base→child so the child's
    own props win.  Cycles are guarded by an id-visited set (defensive; Word
    doesn't emit cyclic basedOn)."""
    chain = []
    seen = set()
    s = style
    while s is not None and id(s) not in seen:
        seen.add(id(s))
        chain.append(s)
        s = s.base_style
    merged: dict = {}
    for anc in reversed(chain):          # root first, child overrides
        merged.update(_extract_style_props(anc))
    return merged


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _docdefaults(styles) -> dict:
    """Read the document-wide default font/size from `w:docDefaults/rPrDefault`.

    Word stores the body baseline here (not on the Normal style) — the NIEHS
    templates leave Normal font-less and set Times New Roman 12pt in docDefaults.
    Returns {'font': ..., 'font_size': ...} with only the keys actually present.
    """
    out: dict = {}
    el = styles.element  # the <w:styles> lxml element python-docx already parsed
    rpr = el.find(f"{{{_W_NS}}}docDefaults/{{{_W_NS}}}rPrDefault/{{{_W_NS}}}rPr")
    if rpr is None:
        return out
    rfonts = rpr.find(f"{{{_W_NS}}}rFonts")
    if rfonts is not None:
        name = rfonts.get(f"{{{_W_NS}}}ascii")
        if name:
            out["font"] = name
    sz = rpr.find(f"{{{_W_NS}}}sz")
    if sz is not None and sz.get(f"{{{_W_NS}}}val"):
        # w:sz is in half-points.
        out["font_size"] = _emu_to_pt_str(int(sz.get(f"{{{_W_NS}}}val")) / 2 * _EMU_PER_PT)
    return out


def _extract_document(section, styles) -> dict:
    """Read page geometry + base fonts from the first section → the `document` layer."""
    doc_layer: dict = {}
    pw = _emu_to_in_str(section.page_width)
    ph = _emu_to_in_str(section.page_height)
    if pw:
        doc_layer["page_width"] = pw
    if ph:
        doc_layer["page_height"] = ph
    for key, val in (
        ("margin_top", section.top_margin),
        ("margin_bottom", section.bottom_margin),
        ("margin_left", section.left_margin),
        ("margin_right", section.right_margin),
        ("header_distance", section.header_distance),
    ):
        s = _emu_to_in_str(val)
        if s:
            doc_layer[key] = s
    # Base body font: prefer the Normal style, fall back to docDefaults (where
    # Word actually stores the body baseline when Normal is left font-less).
    normal_font = styles["Normal"].font
    dd = _docdefaults(styles)
    default_font = normal_font.name or dd.get("font")
    if default_font:
        doc_layer["default_font"] = default_font
    size = _emu_to_pt_str(normal_font.size) or dd.get("font_size")
    if size:
        doc_layer["default_font_size"] = size
    try:
        header_font = styles["Header"].font
        if header_font.name:
            doc_layer["header_font"] = header_font.name
        hsize = _emu_to_pt_str(header_font.size)
        if hsize:
            doc_layer["header_font_size"] = hsize
    except KeyError:
        pass
    return doc_layer


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_styles(docx_path: str) -> dict:
    """
    Read a Word template into a `styles` mapping (the {defaults, types, document}
    shape the pipeline validates + consumes).

    - `defaults` ← the Normal style (the document-wide baseline).
    - `types`    ← Heading 1/2/3 → the heading-bearing catalog node_types.
    - `document` ← first-section geometry + base body/header fonts.

    Only buckets 1+2 (typography + document geometry).  Bucket-3 component
    geometry (table/ToC/list) is NOT read; a note is logged to stderr so the
    omission is explicit.
    """
    doc = _open_word(docx_path)
    styles = doc.styles

    cfg: dict = {}

    # defaults ← Normal, merged over docDefaults (Word's body baseline lives in
    # docDefaults when Normal is font-less, as in the NIEHS templates).
    normal_props = _extract_style_props(styles["Normal"])
    defaults = {**_docdefaults(styles), **normal_props}
    if defaults:
        cfg["defaults"] = defaults

    # types ← each Heading style, fanned out to the node_types that use it.
    types: dict = {}
    for style_name, node_types in _HEADING_STYLE_TO_TYPES.items():
        try:
            props = _extract_style_props(styles[style_name])
        except KeyError:
            continue
        if not props:
            continue
        for nt in node_types:
            types[nt] = dict(props)
    if types:
        cfg["types"] = types

    # title_page ← the NTP `1-NN` title-page style family, keyed by role.  These
    # styles inherit from Base_Heading/Base_Text, so we resolve the basedOn chain
    # (python-docx does not) via _resolved_style_props.
    title_page: dict = {}
    for style_name, role in _TITLE_PAGE_STYLE_TO_ROLE.items():
        try:
            style = styles[style_name]
        except KeyError:
            continue
        props = _resolved_style_props(style)
        if props:
            title_page[role] = props
    if title_page:
        cfg["title_page"] = title_page

    # document ← first section geometry + base fonts.
    if doc.sections:
        doc_layer = _extract_document(doc.sections[0], styles)
        if doc_layer:
            cfg["document"] = doc_layer

    _log_skipped(doc)
    return cfg


def _log_skipped(doc) -> None:
    """Emit a stderr note naming the bucket-3 component geometry we did NOT read,
    so 'not extracted' is visible rather than a silent gap."""
    notes = []
    if doc.tables:
        notes.append(f"{len(doc.tables)} table(s): column widths / cell margins / borders")
    # ToC dot-leaders + list numbering live in styles we don't map.
    notes.append("ToC tab-stops (dot leaders), list indent/numbering")
    print(
        "docx_style_extract: extracted typography + document geometry only. "
        "NOT extracted (per-surface component geometry): " + "; ".join(notes),
        file=sys.stderr,
    )


# ---------------------------------------------------------------------------
# PDF / converter contamination detection
# ---------------------------------------------------------------------------
# Documents round-tripped through PDF import (Acrobat/LibreOffice) accrete
# auto-generated style DEFINITIONS that no content uses — dead library cruft
# that would poison a naive "read every style" extraction.  The NTP "Publication
# Version Draft" examples carry ~150 such extras vs. the ~384 in the clean
# _Final files (see the docx-governance memory).  We name the fingerprint so the
# coverage diagnostic can WARN, and so a vocabulary generated with used_only can
# report how much it excluded rather than silently dropping it.

# Auto-generated converter names: Acrobat clusters observed formatting into
# CM<n> ("character map") and Pa<n> ("paragraph") styles — no human authors these.
_CONVERTER_NAME_RE = re.compile(r"^(CM|Pa)\d+$")

# Word appends a digit when a pasted/merged style collides with an existing one;
# cascading suffixes (No List1 → No List11) mark repeated paste cycles.  We flag a
# name that is an existing name + a trailing digit only when the base also exists.
_TRAILING_DIGIT_RE = re.compile(r"^(.*?)(\d+)$")


def _classify_contamination(style_names: set[str]) -> dict:
    """Partition a template's style NAMES into contamination classes.

    Returns {converter: [...], numbered_dup: [...]} — the two machine-detectable
    PDF/paste-import fingerprints.  ``converter`` is the strong signal (CM##/Pa#,
    unambiguously Acrobat); ``numbered_dup`` is a name that equals an existing
    style plus a trailing digit (a paste-collision twin like 'No List11')."""
    converter = sorted(n for n in style_names if _CONVERTER_NAME_RE.match(n))
    numbered_dup = []
    for n in style_names:
        m = _TRAILING_DIGIT_RE.match(n)
        if not m:
            continue
        base = m.group(1).rstrip()
        # A collision twin only if stripping the trailing digit yields an
        # existing style name (and the base isn't itself a converter artifact).
        if base and base in style_names and not _CONVERTER_NAME_RE.match(n):
            numbered_dup.append(n)
    return {"converter": converter, "numbered_dup": sorted(numbered_dup)}


def coverage_report(docx_path: str) -> dict:
    """Report which EXPECTED styles a template actually contains, without extracting.

    The extractor maps a fixed set of Word style NAMES onto the vocabulary
    (Normal → defaults, Heading 1/2/3 → types, the NTP ``1-NN`` family →
    title_page roles).  When a name is absent it is silently skipped
    (``except KeyError: continue``), so extraction "succeeds" while emitting a thin
    config — the failure mode when a real template names things differently than
    the reverse-engineered one we built the maps against.

    This turns that silent gap into an explicit read: for each expected name it
    reports found/missing and (when found) whether the style resolves to any
    props; it also lists the template's OTHER paragraph styles that carry real
    formatting — mapping candidates the current maps don't reach.  Read-only; it
    does not build or return a styles config.

    Returns::

        {
          "expected": [
            {"style": "Normal", "target": "defaults", "present": True, "has_props": True},
            {"style": "1-03_Report_Title", "target": "title_page:report_title",
             "present": False, "has_props": False},
            ...
          ],
          "found": int, "missing": int, "total": int,
          "unmapped_with_props": ["SomeCustomHeading", ...],
          "contamination": {
            "converter": ["CM14", "Pa2", ...],       # Acrobat PDF-import artifacts
            "numbered_dup": ["No List11", ...],       # paste-collision twins
            "total": int, "used_in_body": [...],      # cruft actually applied (rare)
          },
        }
    """
    doc = _open_word(docx_path)
    styles = doc.styles

    # (style name, human target) for every name the extractor tries to read.
    expected: list[tuple[str, str]] = [
        ("Normal", "defaults"),
        ("Header", "document:header_font"),
    ]
    for style_name in _HEADING_STYLE_TO_TYPES:
        expected.append((style_name, "types"))
    for style_name, role in _TITLE_PAGE_STYLE_TO_ROLE.items():
        expected.append((style_name, f"title_page:{role}"))

    def _lookup(name: str):
        try:
            return styles[name]
        except KeyError:
            return None

    rows: list[dict] = []
    mapped_names = set()
    for name, target in expected:
        mapped_names.add(name)
        style = _lookup(name)
        present = style is not None
        has_props = bool(_resolved_style_props(style)) if present else False
        rows.append({
            "style": name, "target": target,
            "present": present, "has_props": has_props,
        })

    # Other paragraph styles carrying real formatting = mapping candidates the
    # maps don't reach.  Restrict to paragraph styles (WD_STYLE_TYPE.PARAGRAPH == 1)
    # with resolvable props, excluding the ones we already expect.
    unmapped: list[str] = []
    for style in styles:
        name = getattr(style, "name", None)
        if not name or name in mapped_names:
            continue
        if int(getattr(style, "type", 0) or 0) != 1:  # paragraph styles only
            continue
        if _resolved_style_props(style):
            unmapped.append(name)

    # Contamination: PDF/paste-import style DEFINITIONS.  The strongest signal is
    # that they are UNUSED — Word never GCs styles, so import cruft accretes as
    # dead definitions no paragraph references.  We report the fingerprint classes
    # and, for each, how many are actually applied in the body (used cruft is a
    # real concern; unused cruft is harmless and excluded by used_only extraction).
    all_para_names = {
        s.name for s in styles
        if getattr(s, "name", None) and int(getattr(s, "type", 0) or 0) == 1
    }
    classes = _classify_contamination(all_para_names)
    used_names = {
        p.style.name for p in doc.paragraphs if p.style is not None
    }
    contamination = {
        "converter": classes["converter"],
        "numbered_dup": classes["numbered_dup"],
        "total": len(classes["converter"]) + len(classes["numbered_dup"]),
        "used_in_body": sorted(
            (set(classes["converter"]) | set(classes["numbered_dup"])) & used_names
        ),
    }

    found = sum(1 for r in rows if r["present"])
    return {
        "expected": rows,
        "found": found,
        "missing": len(rows) - found,
        "total": len(rows),
        "unmapped_with_props": sorted(unmapped),
        "contamination": contamination,
    }


def format_coverage(report: dict) -> str:
    """Render coverage_report() as a human-readable text block for the CLI."""
    lines: list[str] = []
    lines.append(
        f"Template style coverage: {report['found']}/{report['total']} "
        f"expected styles present ({report['missing']} missing)."
    )
    lines.append("")
    for r in report["expected"]:
        if not r["present"]:
            mark = "MISSING"
        elif r["has_props"]:
            mark = "found  "
        else:
            mark = "empty  "  # present but resolves to no props (inherits all)
        lines.append(f"  [{mark}] {r['style']:<22} -> {r['target']}")
    unmapped = report["unmapped_with_props"]
    if unmapped:
        lines.append("")
        lines.append(
            f"Unmapped paragraph styles with formatting "
            f"({len(unmapped)}) — mapping candidates:"
        )
        for name in unmapped:
            lines.append(f"    {name}")

    contamination = report.get("contamination") or {}
    if contamination.get("total"):
        conv = contamination["converter"]
        dup = contamination["numbered_dup"]
        used = contamination["used_in_body"]
        lines.append("")
        lines.append(
            f"WARNING: {contamination['total']} likely PDF/paste-import style "
            f"artifact(s) detected — this document looks converter-contaminated:"
        )
        if conv:
            lines.append(
                f"    {len(conv)} Acrobat converter style(s) (CM##/Pa#): "
                + ", ".join(conv[:10]) + (" ..." if len(conv) > 10 else "")
            )
        if dup:
            lines.append(
                f"    {len(dup)} paste-collision duplicate(s): "
                + ", ".join(dup[:10]) + (" ..." if len(dup) > 10 else "")
            )
        if used:
            lines.append(
                f"    {len(used)} of these are APPLIED in the body (real risk): "
                + ", ".join(used[:10])
            )
        else:
            lines.append(
                "    none are applied in the body — dead library cruft, excluded "
                "by used_only extraction (harmless)."
            )
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Vocabulary generation — a Word template's named-style graph → a vocabulary.yaml
# ---------------------------------------------------------------------------
# The inverse-plus of extract_styles: instead of flattening a handful of styles
# into the {defaults, types, title_page} styling config, this walks the FULL
# paragraph-style graph (every used style + its basedOn parent) and emits the
# descriptive-markup VOCABULARY (vocabulary.py): one type record per style, with
#   specializes ← the basedOn parent (slugified to a type name),
#   style       ← the style's OWN delta (its props minus what the parent sets),
#   bind.docx   ← the real Word style name (so the docx surface round-trips).
# The specialization GRAPH is preserved (deltas, not resolved absolutes), so the
# vocabulary carries Word's inheritance rather than freezing it.

_SLUG_RE = re.compile(r"[^a-z0-9]+")


def _slug(word_style_name: str) -> str:
    """A Word style name → a vocabulary type name (snake_case slug).

    '1-03_Report_Title' → 'report_title_1_03'? No — we keep it stable and
    reversible-ish: lowercase, non-alnum runs → '_', strip edges.  The NTP names
    already encode the role ('Report_Title'); we drop the leading numeric family
    prefix ('1-03_') so the type name reads as the ROLE, and keep bind.docx as the
    exact original for the round-trip."""
    name = word_style_name.strip()
    # Drop a leading NTP family prefix like "1-03_", "0-29a_", "4-10a_".
    name = re.sub(r"^[0-9]+-[0-9a-z]+_", "", name, flags=re.IGNORECASE)
    slug = _SLUG_RE.sub("_", name.lower()).strip("_")
    return slug or _SLUG_RE.sub("_", word_style_name.lower()).strip("_")


def _own_style_delta(style, parent_props: dict) -> dict:
    """A style's OWN props minus what its parent already sets (the specialization
    delta).  A key whose value equals the parent's is inherited, so it is dropped
    — only genuine overrides remain, mirroring how Word stores basedOn styles."""
    own = _extract_style_props(style)
    return {k: v for k, v in own.items() if parent_props.get(k) != v}


# Where an NTP root style attaches into the neutral base vocabulary (base.yaml).
# These are the only cross-vocabulary edges: the domain roots specialize a
# medium-neutral base type so the whole graph resolves through one chain.
_NTP_ROOT_TO_BASE = {
    "Base_Text": "block",
    "Base_Heading": "heading",
    "Normal": "text",
}

# Styles to include in the vocabulary even when the example body does not APPLY
# them (used_only would otherwise skip them).  These are semantic roles our
# generator emits or will emit, that the NTP template DEFINES but its example
# bodies happen not to exercise — so they are real, wanted vocabulary, not the
# unused-cruft used_only is meant to drop.  Present + identical across all five
# examples (verified): the title subtitle/number lines and the table-body cell
# style.  basedOn ancestors of these are pulled in transitively.
_ALWAYS_INCLUDE = (
    "1-04_Report_Subtitle",   # report_subtitle role (title page)
    "1-05a_Report-Number",    # report_number role (title page)
    "0-29_Table_Body",        # table_body_cell role (table cells)
    "Fig_Title",              # fig_title role (figure, ADR-0012)
    "3-11_Fig_Alt_Text",      # fig_alt_text role (figure accessibility)
    "1-26_Logo_Graphic",      # logo_graphic role (figure subtype=logo)
    "NTP Publisher Location", # publisher_location role (title-page location line)
)


def generate_vocabulary(
    docx_path: str,
    name: str = "ntp-report",
    extends: str = "base",
    used_only: bool = True,
    always_include: "tuple[str, ...]" = _ALWAYS_INCLUDE,
) -> dict:
    """Walk a Word template's paragraph-style graph → a vocabulary dict (the shape
    vocabulary.load_vocabulary consumes: {vocabulary, extends, types}).

    Every emitted PARAGRAPH style becomes a type; ``specializes`` is its basedOn
    parent (slugified); ``style`` is its delta vs the parent's RESOLVED props;
    ``bind`` pins the docx name.  A domain ROOT style (Base_Text/Base_Heading/
    Normal — no basedOn within the set) specializes a neutral base type via
    ``_NTP_ROOT_TO_BASE`` so the graph resolves through base.yaml.

    ``used_only`` (default) restricts output to styles actually USED in the
    document body PLUS their basedOn ancestors — the real report vocabulary
    (~50), not every built-in Word style Word carries (~230).  Set False to emit
    the full paragraph-style catalog.

    ``always_include`` names styles to keep even when the body does not apply them
    (their basedOn ancestors come along transitively).  These are template-DEFINED
    semantic roles our generator emits that the example bodies happen not to
    exercise — real vocabulary, not the unused cruft ``used_only`` drops.  A name
    absent from this template is silently skipped."""
    doc = _open_word(docx_path)
    styles = doc.styles

    para_styles = [s for s in styles if int(getattr(s, "type", 0) or 0) == 1]
    by_name = {s.name: s for s in para_styles}

    if used_only:
        keep: set[str] = set()
        used = {p.style.name for p in doc.paragraphs if p.style is not None}
        # Seed with body-applied styles AND the always-include roles, then pull in
        # every seed's basedOn ancestors so no chain dangles.
        seeds = set(used) | {n for n in always_include if n in by_name}
        for uname in seeds:
            s = by_name.get(uname)
            while s is not None and s.name not in keep:  # include basedOn ancestors
                keep.add(s.name)
                s = s.base_style
        emit = [s for s in para_styles if s.name in keep]
    else:
        emit = para_styles

    types: dict = {}
    for s in emit:
        parent = s.base_style
        parent_props = _resolved_style_props(parent) if parent is not None else {}
        delta = _own_style_delta(s, parent_props)
        if parent is not None and parent.name in by_name:
            specializes = _slug(parent.name)
        else:
            specializes = _NTP_ROOT_TO_BASE.get(s.name)  # domain root → neutral base type
        entry: dict = {}
        if specializes:
            entry["specializes"] = specializes
        if delta:
            entry["style"] = delta
        entry["bind"] = {"docx": s.name}
        types[_slug(s.name)] = entry

    return {"vocabulary": name, "extends": extends, "types": types}


def vocabulary_to_yaml(vocab: dict) -> str:
    """Serialize a generated vocabulary dict as a YAML document."""
    import yaml
    return yaml.safe_dump(vocab, sort_keys=False, default_flow_style=False)


def to_yaml(cfg: dict) -> str:
    """Serialize an extracted config as a `styles:`-wrapped YAML document."""
    import yaml
    return yaml.safe_dump({"styles": cfg}, sort_keys=False, default_flow_style=False)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _main(argv: list[str]) -> int:
    import argparse
    parser = argparse.ArgumentParser(
        prog="docx_style_extract",
        description="Extract a styles.yaml (typography + document geometry) from a Word template.",
    )
    parser.add_argument("docx", help="path to the .docx template")
    parser.add_argument("-o", "--out", help="write YAML here (default: stdout)")
    parser.add_argument(
        "--coverage", action="store_true",
        help="report found/missing expected styles instead of extracting (a "
             "day-one read on how well a new template matches the style maps)",
    )
    parser.add_argument(
        "--emit-vocabulary", action="store_true",
        help="emit a vocabulary.yaml (the full named-style graph as semantic "
             "types with specializes/style/bind) instead of the styling config",
    )
    parser.add_argument(
        "--vocab-name", default="ntp-report",
        help="name for the emitted vocabulary (default: ntp-report)",
    )
    parser.add_argument(
        "--vocab-extends", default="base",
        help="parent vocabulary the emitted one extends (default: base)",
    )
    args = parser.parse_args(argv)

    if args.coverage:
        print(format_coverage(coverage_report(args.docx)))
        return 0

    if args.emit_vocabulary:
        vocab = generate_vocabulary(
            args.docx, name=args.vocab_name, extends=args.vocab_extends
        )
        text = vocabulary_to_yaml(vocab)
        if args.out:
            with open(args.out, "w", encoding="utf-8") as f:
                f.write(text)
            print(f"wrote {args.out}", file=sys.stderr)
        else:
            sys.stdout.write(text)
        return 0

    cfg = extract_styles(args.docx)
    yaml_text = to_yaml(cfg)
    if args.out:
        with open(args.out, "w", encoding="utf-8") as f:
            f.write(yaml_text)
        print(f"wrote {args.out}", file=sys.stderr)
    else:
        sys.stdout.write(yaml_text)
    return 0


if __name__ == "__main__":
    raise SystemExit(_main(sys.argv[1:]))
