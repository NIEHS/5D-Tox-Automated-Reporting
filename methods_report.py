"""
methods_report.py — backward-compatible facade + DOCX assembly.

The original ~3200-line monolith was extracted into focused submodules
across the methods_* and abstract_* commits on this branch:

  methods_models.py        SUBSECTION_SKELETON + MethodsContext +
                           MethodsSection + MethodsReport dataclasses
  methods_extract.py       extract_methods_context + the bm2 / PK /
                           biosampling / genomics-assay / sample-count
                           extractors that feed it
  methods_prompt.py        build_methods_prompt + build_subsection_skeleton
                           + per-subsection LLM guideline composer
  methods_table1.py        build_table1_data — the only programmatic
                           table in the Methods section
  narrative_helpers.py     cross-cutting formatters used by 4+ narrative
                           sections (_join_oxford, _format_dose_value,
                           _is_reliable_bmd, _picks_above_lle, etc.)
  abstract_methods.py      Abstract → Methods paragraph (deterministic)
  abstract_apical.py       body-Results per-platform narrative +
                           Abstract → Results apical paragraph
  abstract_genomics.py     Abstract → Results genomics paragraph +
                           sentence builders
  abstract_pk.py           Abstract → Results pharmacokinetics paragraph
  gene_bodies.py           body Results: Gene Set BMD + Gene BMD analysis
                           intro/findings paragraphs
  abstract_summary.py      Abstract → Summary + Abstract → Results
                           aggregator (the only abstract_* module that
                           imports from sibling abstract_* modules)

What still lives here:

  - re-exports of every public + private name above, so existing
    importers don't need to change
  - add_methods_to_doc + _add_methods_table: the python-docx
    assembly that turns a MethodsReport into rendered Word output.
    They live here because they have no callers outside the DOCX
    export path; if a future commit splits the docx renderer out,
    these go with it.

External importers preserved through this shim:

  - genomics_narratives.py  build_gene_set_body_{intro,findings},
                            build_gene_body_{intro,findings}
  - llm_routes.py           extract_methods_context + build_methods_prompt
                            + build_subsection_skeleton + build_table1_data
                            + MethodsReport + MethodsSection
  - process_integrated.py   build_apical_bmd_summary_narrative,
                            same MethodsReport / extract_methods_context
                            set as llm_routes
  - report_data.py           MethodsContext + build_abstract_methods +
                            build_abstract_results + build_abstract_summary
  - processing_helpers.py   _is_anomalous_bmd (lazy import inside a
                            helper)

A future cleanup pass could point each consumer at the new modules
directly and delete this shim, but that's not part of this split.
"""

# ---------------------------------------------------------------------------
# Re-exports: dataclasses + heading skeleton
# ---------------------------------------------------------------------------
from methods_models import (
    SUBSECTION_SKELETON,
    MethodsContext,
    MethodsSection,
    MethodsReport,
)

# ---------------------------------------------------------------------------
# Re-exports: cross-cutting narrative helpers
# ---------------------------------------------------------------------------
from narrative_helpers import (
    _DIRECTION_WORDS,
    ANOMALY_RATIO_THRESHOLD,
    _is_reliable_bmd,
    _is_anomalous_bmd,
    _format_dose_value,
    _format_rat_gene_symbol,
    _stat_display_name,
    _join_oxford,
    _format_dose_list,
    _format_organ_list,
    _format_organ_phrase,
    _normalize_organ_name,
    _format_paired_bmd_pairs,
    _picks_above_lle,
)

# ---------------------------------------------------------------------------
# Re-exports: methods context extractor
# ---------------------------------------------------------------------------
from methods_extract import (
    _parse_bm2_analysis_info,
    _collect_bm2_analysis_metadata,
    extract_methods_context,
    _extract_biosampling_doses,
    _extract_pk_data,
    _extract_genomics_assay,
    _build_genomics_sample_counts,
)

# ---------------------------------------------------------------------------
# Re-exports: LLM prompt assembly
# ---------------------------------------------------------------------------
from methods_prompt import (
    build_subsection_skeleton,
    build_methods_prompt,
    _build_subsection_guidelines,
)

# ---------------------------------------------------------------------------
# Re-exports: Table 1 builder
# ---------------------------------------------------------------------------
from methods_table1 import build_table1_data

# ---------------------------------------------------------------------------
# Re-exports: narrative builders (apical, genomics, pk, gene bodies, summary)
# ---------------------------------------------------------------------------
from abstract_methods import build_abstract_methods
from abstract_apical import (
    _normalize_endpoint_name,
    _format_endpoint_phrase,
    _format_bmd_pair,
    build_apical_bmd_summary_narrative,
    build_abstract_results_apical,
)
from abstract_genomics import (
    _build_gene_sets_sentence,
    _build_top_genes_sentence,
    build_abstract_results_genomics,
)
from abstract_pk import build_abstract_results_pk
from gene_bodies import (
    build_gene_set_body_intro,
    build_gene_set_body_findings,
    build_gene_body_intro,
    build_gene_body_findings,
)
from abstract_summary import (
    build_abstract_summary,
    build_abstract_results,
)


# ---------------------------------------------------------------------------
# DOCX generation: add structured M&M to a python-docx Document
# ---------------------------------------------------------------------------
# These two functions remain inline because nothing else in the codebase
# uses them — they only run from the /api/export-docx path that calls
# add_methods_to_doc to splice the Methods section into the assembled
# Word document.  Keeping them here avoids a single-caller module.

def add_methods_to_doc(
    doc,
    methods_report: MethodsReport,
    start_table_num: int = 1,
) -> int:
    """
    Add the structured Materials and Methods section to a python-docx Document.

    Generates the full NIEHS-style M&M with hierarchical headings (H2, H3, H4),
    prose paragraphs per subsection, and Table 1 (genomics sample counts).

    Follows the same DOCX formatting conventions as add_animal_report_to_doc():
    - Calibri 11pt for body text
    - Calibri 9pt for table cells and captions
    - "Light Shading Accent 1" table style
    - Sequential table numbering

    Args:
        doc:             python-docx Document object.
        methods_report:  MethodsReport with sections and context populated.
        start_table_num: Table number to start from (for sequential numbering).

    Returns:
        Next table number (for subsequent sections to continue from).
    """
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table_num = start_table_num

    # --- Top-level heading ---
    doc.add_heading("Materials and Methods", level=2)

    # --- Render each subsection ---
    for section in methods_report.sections:
        # Add heading at the appropriate level
        doc.add_heading(section.heading, level=section.level)

        # Add prose paragraphs
        for para_text in section.paragraphs:
            if not para_text.strip():
                continue
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(6)

        # Add table if present (e.g. Table 1)
        if section.table:
            table_num = _add_methods_table(
                doc, section.table, table_num,
            )

    # --- Add Table 1 at the end if genomics data exists ---
    # Table 1 is appended after all prose subsections, matching the
    # NIEHS report layout where it follows the Data Analysis section.
    table1_data = build_table1_data(methods_report.context)
    if table1_data:
        table_num = _add_methods_table(doc, table1_data, table_num)

    return table_num


def _add_methods_table(doc, table_data: dict, table_num: int) -> int:
    """
    Add a formatted table to the DOCX document.

    Handles the caption-above-table pattern, bold sex-header rows,
    and footnotes below the table.

    Args:
        doc:        python-docx Document.
        table_data: Dict with caption, headers, rows, footnotes.
        table_num:  Current table number for caption.

    Returns:
        Next table number.
    """
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT

    headers = table_data["headers"]
    rows = table_data["rows"]
    caption_text = table_data.get("caption", "")
    footnotes = table_data.get("footnotes", [])

    n_cols = len(headers)
    n_rows = len(rows)

    # Create the table
    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Caption paragraph (inserted before the table)
    caption = doc.add_paragraph()
    run = caption.add_run(f"Table {table_num}. {caption_text}")
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.italic = True
    caption.paragraph_format.space_after = Pt(4)
    # Move caption before the table element in the document XML
    table._element.addprevious(caption._element)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = "Calibri"

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[1 + row_idx]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            # Check for bold marker (sex header rows: "**Male**")
            is_bold = val.startswith("**") and val.endswith("**")
            clean_val = val.strip("*").strip() if is_bold else val.strip()
            cell.text = clean_val
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Calibri"
                    if is_bold:
                        r.bold = True

    # Footnotes below the table
    for fn in footnotes:
        p = doc.add_paragraph()
        run = p.add_run(fn)
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        run.italic = True
        p.paragraph_format.space_after = Pt(2)

    return table_num + 1
