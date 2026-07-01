"""
Unit tests for freeform authored-content components (freeform-page /
freeform-block).

Three layers are pinned here:

  1. The resolver (freeform_content.resolve_freeform): single-source latex/html,
     dual-source mapping, docx → both surfaces, content_file read relative to a
     base dir, and the pending-note helper.
  2. The instantiator validation (document_template._validate_freeform_entry):
     exactly-one-source, representation required/valid, docx-needs-file, dual-
     source-forbids-representation, and the stray-binding-on-non-freeform guard.
  3. The two renderer handlers: a freeform-page forces a page break (LaTeX
     \\clearpage / HTML break-before:page) and a freeform-block does not; the
     foreign surface shows a pending note while the native surface emits the
     authored markup verbatim.

The docx fixture is built in-test with python-docx (no binary committed).
"""

import pathlib

import pytest

import freeform_content as fc
from document_node import DocNode
from document_template import instantiate

TEMPLATES_DIR = pathlib.Path(__file__).resolve().parents[2] / "templates"


# ---------------------------------------------------------------------------
# Resolver — single source, dual source
# ---------------------------------------------------------------------------

def test_resolve_latex_is_native_to_latex_only():
    out = fc.resolve_freeform(r"\textbf{hi}", None, "latex", base_dir=TEMPLATES_DIR)
    assert out == {"latex": r"\textbf{hi}", "html": None}


def test_resolve_html_is_native_to_html_only():
    out = fc.resolve_freeform("<b>hi</b>", None, "html", base_dir=TEMPLATES_DIR)
    assert out == {"latex": None, "html": "<b>hi</b>"}


def test_resolve_dual_source_keeps_each_surface_verbatim():
    content = {"latex": r"\emph{x}", "html": "<em>x</em>"}
    out = fc.resolve_freeform(content, None, None, base_dir=TEMPLATES_DIR)
    assert out == {"latex": r"\emph{x}", "html": "<em>x</em>"}


def test_resolve_dual_source_missing_key_is_none_on_that_surface():
    out = fc.resolve_freeform({"html": "<em>x</em>"}, None, None, base_dir=TEMPLATES_DIR)
    assert out == {"latex": None, "html": "<em>x</em>"}


def test_resolve_content_file_read_relative_to_base_dir(tmp_path):
    (tmp_path / "frag.tex").write_text(r"\section*{Hi}", encoding="utf-8")
    out = fc.resolve_freeform(None, "frag.tex", "latex", base_dir=tmp_path)
    assert out == {"latex": r"\section*{Hi}", "html": None}


def test_resolve_missing_content_file_raises(tmp_path):
    with pytest.raises(ValueError, match="content_file not found"):
        fc.resolve_freeform(None, "nope.html", "html", base_dir=tmp_path)


def test_resolve_dual_file_reads_each_surface_from_its_own_file(tmp_path):
    (tmp_path / "a.tex").write_text(r"\emph{tex}", encoding="utf-8")
    (tmp_path / "a.html").write_text("<em>html</em>", encoding="utf-8")
    out = fc.resolve_freeform(
        None, {"latex": "a.tex", "html": "a.html"}, None, base_dir=tmp_path
    )
    assert out == {"latex": r"\emph{tex}", "html": "<em>html</em>"}


def test_resolve_dual_file_missing_key_is_none_on_that_surface(tmp_path):
    (tmp_path / "only.html").write_text("<em>html</em>", encoding="utf-8")
    out = fc.resolve_freeform(
        None, {"html": "only.html"}, None, base_dir=tmp_path
    )
    assert out == {"latex": None, "html": "<em>html</em>"}


# ---------------------------------------------------------------------------
# Resolver — docx → both surfaces (fixture built in-test)
# ---------------------------------------------------------------------------

def _build_docx_fixture(path: pathlib.Path) -> None:
    """A small .docx exercising every block kind in the documented subset:
    a heading, a paragraph with a bold run, a 2-item bullet list, a 2x2 table."""
    import docx

    document = docx.Document()
    document.add_heading("Findings", level=1)
    p = document.add_paragraph("This is ")
    p.add_run("bold").bold = True
    document.add_paragraph("first bullet", style="List Bullet")
    document.add_paragraph("second bullet", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    document.save(str(path))


def test_resolve_docx_yields_both_surfaces(tmp_path):
    _build_docx_fixture(tmp_path / "doc.docx")
    out = fc.resolve_freeform(None, "doc.docx", "docx", base_dir=tmp_path)

    # LaTeX surface: section heading, bold run, itemize list, tabular.
    latex = out["latex"]
    assert r"\section*{Findings}" in latex
    assert r"\textbf{bold}" in latex
    assert r"\begin{itemize}" in latex
    assert r"\begin{tabular}" in latex
    assert "A & B" in latex

    # HTML surface: h2 (page title owns h1), strong run, ul list, table.
    html = out["html"]
    assert "<h2>Findings</h2>" in html
    assert "<strong>bold</strong>" in html
    assert "<ul><li>first bullet</li><li>second bullet</li></ul>" in html
    assert '<table class="freeform-table">' in html
    assert "<td>A</td>" in html


def test_resolve_docx_without_file_raises(tmp_path):
    with pytest.raises(ValueError, match="docx.*requires content_file"):
        fc.resolve_freeform(None, None, "docx", base_dir=tmp_path)


# ---------------------------------------------------------------------------
# Pending note
# ---------------------------------------------------------------------------

def test_pending_note_points_at_the_other_surface():
    # latex content viewed on the HTML surface points at the LaTeX export.
    assert "Overleaf/LaTeX export" in fc.pending_note("latex", "html")
    # html content viewed on the LaTeX surface points at the HTML preview.
    assert "HTML preview" in fc.pending_note("html", "latex")


# ---------------------------------------------------------------------------
# Instantiator validation
# ---------------------------------------------------------------------------

def _page(**kw) -> dict:
    base = {"id": "ff", "type": "freeform-page", "title": "FF"}
    base.update(kw)
    return base


def test_validation_rejects_both_content_and_content_file():
    with pytest.raises(ValueError, match="exactly one"):
        instantiate([_page(content="x", content_file="f.tex", representation="latex")])


def test_validation_rejects_neither_content_nor_file():
    with pytest.raises(ValueError, match="requires `content` or `content_file`"):
        instantiate([_page(representation="latex")])


def test_validation_rejects_missing_representation_single_source():
    with pytest.raises(ValueError, match="`representation` is required"):
        instantiate([_page(content="x")])


def test_validation_rejects_bad_representation():
    with pytest.raises(ValueError, match="must be one of"):
        instantiate([_page(content="x", representation="markdown")])


def test_validation_rejects_docx_without_file():
    with pytest.raises(ValueError, match="docx.*requires"):
        instantiate([_page(content="x", representation="docx")])


def test_validation_rejects_representation_with_dual_source():
    with pytest.raises(ValueError, match="not allowed with"):
        instantiate([_page(content={"latex": "a", "html": "b"}, representation="latex")])


def test_validation_rejects_dual_source_without_latex_or_html_key():
    with pytest.raises(ValueError, match="latex.*html"):
        instantiate([_page(content={"foo": "bar"})])


def test_validation_accepts_dual_file_mapping(tmp_path):
    # A dict content_file with latex+html keys is the dual-file form; it must
    # pass validation (no representation, at least one surface key present).
    tree = instantiate(
        [_page(content_file={"latex": "freeform/appendix-e.tex",
                             "html": "freeform/appendix-e.html"})]
    )
    node = tree[0]
    assert node.resolved_content["latex"]
    assert node.resolved_content["html"]


def test_validation_rejects_dual_file_without_latex_or_html_key():
    with pytest.raises(ValueError, match="latex.*html"):
        instantiate([_page(content_file={"foo": "bar.tex"})])


def test_validation_rejects_representation_with_dual_file():
    with pytest.raises(ValueError, match="not allowed with"):
        instantiate([_page(content_file={"latex": "a.tex"}, representation="latex")])


def test_validation_rejects_content_on_non_freeform_type():
    with pytest.raises(ValueError, match="only valid on"):
        instantiate([{"id": "n", "type": "narrative", "title": "N",
                      "data_key": "d", "content": "x", "representation": "latex"}])


def test_instantiate_resolves_content_on_the_node():
    tree = instantiate([_page(content=r"\emph{hi}", representation="latex")])
    node = tree[0]
    assert node.resolved_content == {"latex": r"\emph{hi}", "html": None}


# ---------------------------------------------------------------------------
# Renderer handlers — LaTeX
# ---------------------------------------------------------------------------

def _node(node_type, **kw) -> DocNode:
    return DocNode(id="ff", title=kw.pop("title", "Title"), node_type=node_type,
                   level=kw.pop("level", 1), **kw)


def test_latex_page_forces_clearpage_and_emits_authored_latex():
    import latex_generator as lg

    node = _node("freeform-page", representation="latex",
                 resolved_content={"latex": r"\textbf{body}", "html": None})
    out = lg._render_freeform_page(node, {})
    assert r"\clearpage" in out
    assert r"\textbf{body}" in out


def test_latex_block_omits_clearpage():
    import latex_generator as lg

    node = _node("freeform-block", representation="latex",
                 resolved_content={"latex": r"\textbf{body}", "html": None})
    out = lg._render_freeform_block(node, {})
    assert r"\clearpage" not in out
    assert r"\textbf{body}" in out


def test_latex_shows_pending_note_for_html_only_content():
    import latex_generator as lg

    node = _node("freeform-page", representation="html",
                 resolved_content={"latex": None, "html": "<b>x</b>"})
    out = lg._render_freeform_page(node, {})
    assert "HTML preview" in out          # the foreign-surface note
    assert "<b>x</b>" not in out          # authored html never leaks into LaTeX


# ---------------------------------------------------------------------------
# Renderer handlers — HTML
# ---------------------------------------------------------------------------

def test_html_page_forces_break_and_emits_authored_html():
    import html_generator as hg

    node = _node("freeform-page", representation="html",
                 resolved_content={"latex": None, "html": "<p>body</p>"})
    out = hg._render_freeform_page(node, {})
    assert "break-before:page" in out
    assert "<p>body</p>" in out


def test_html_block_omits_break():
    import html_generator as hg

    node = _node("freeform-block", representation="html",
                 resolved_content={"latex": None, "html": "<p>body</p>"})
    out = hg._render_freeform_block(node, {})
    assert "break-before:page" not in out
    assert 'class="freeform-block"' in out
    assert "<p>body</p>" in out


def test_html_shows_pending_note_for_latex_only_content():
    import html_generator as hg

    node = _node("freeform-page", representation="latex",
                 resolved_content={"latex": r"\textbf{x}", "html": None})
    out = hg._render_freeform_page(node, {})
    assert "Overleaf/LaTeX export" in out
    assert r"\textbf{x}" not in out
