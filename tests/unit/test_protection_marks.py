r"""
test_protection_marks.py — the render-channel half of the ADR-0014 per-node
protection mark (step 5).

The workflow's human-facing guard (workflow.guard.human_guard) is the intensity
of a visual "protected" mark.  Step 5 SURFACES a pre-resolved per-node guard
level so every output surface (LaTeX, HTML, docx, JATS/BITS) can draw it; it does
NOT compute the level from facts (that is a later step).  The level map is
threaded on the report data under ``data["protection"]`` — a plain
``{node_id -> GuardLevel}`` dict keyed by the globally-unique DocNode.id.

What this proves, per surface:
  (a) SAFETY — an absent protection key AND an explicit empty map both render
      byte-identical to the baseline (the critical no-op property, mirroring the
      layout-style / ADR-0005 override overlays).
  (b) PRESENCE — with ``{"<real node id>": GuardLevel.GUARDED}`` the surface's
      protection mark appears for THAT node and not for an unrelated one.

Plus resolver unit tests for render_common.resolve_protection's coercion
(GuardLevel / int / str name / stringified int / junk → OPEN).

Node ids come from the real DOCUMENT_TREE (via scaffold_report_data, the same
fixture the html/latex/docx/jats renderer tests consume) so the keys are real.
"""

from io import BytesIO

import pytest
from docx import Document
from docx.oxml.ns import qn

from rendering.report_data import scaffold_report_data
from rendering.latex_generator import generate_latex
from rendering.html_generator import generate_html
from rendering.docx_generator import generate_docx
from rendering.jats_generator import generate_jats, generate_bits
from rendering.render_common import (
    resolve_protection,
    is_protected,
)
from workflow.guard import GuardLevel


# A guarded node and an unrelated one, both real ids in DOCUMENT_TREE (see the
# probe in the renderer tests: "background" is a body narrative, "toc" is front
# matter — distinct, always rendered on every surface).
GUARDED_ID = "background"
OTHER_ID = "toc"


@pytest.fixture(scope="module")
def scaffold() -> dict:
    """Pure scaffold data — the same fixture the renderer tests use."""
    return scaffold_report_data(
        chemical_name="Perfluorohexanesulfonamide",
        casrn="41997-13-1",
        dtxsid="DTXSID50469320",
    )


@pytest.fixture(scope="module")
def guarded(scaffold) -> dict:
    """Scaffold + a protection map marking one node GUARDED."""
    return {**scaffold, "protection": {GUARDED_ID: GuardLevel.GUARDED}}


# ---------------------------------------------------------------------------
# Resolver (render_common.resolve_protection) — pure lookup + coercion
# ---------------------------------------------------------------------------

def test_resolve_protection_absent_and_empty_are_open():
    assert resolve_protection("anything", {}) is GuardLevel.OPEN
    assert resolve_protection("anything", {"protection": {}}) is GuardLevel.OPEN
    assert resolve_protection("anything", {"protection": None}) is GuardLevel.OPEN


def test_resolve_protection_missing_node_is_open():
    data = {"protection": {"other": GuardLevel.GUARDED}}
    assert resolve_protection("me", data) is GuardLevel.OPEN


@pytest.mark.parametrize(
    "stored,expected",
    [
        (GuardLevel.GUARDED, GuardLevel.GUARDED),
        (GuardLevel.PUBLISHED, GuardLevel.PUBLISHED),
        (1, GuardLevel.GUARDED),          # int value
        (2, GuardLevel.PUBLISHED),
        ("GUARDED", GuardLevel.GUARDED),  # member name
        ("published", GuardLevel.PUBLISHED),
        ("  Guarded  ", GuardLevel.GUARDED),  # whitespace + case tolerant
        ("1", GuardLevel.GUARDED),        # stringified int (JSON-lossy transport)
        (True, GuardLevel.GUARDED),       # bool flag → protected
    ],
)
def test_resolve_protection_coerces(stored, expected):
    assert resolve_protection("n", {"protection": {"n": stored}}) is expected


@pytest.mark.parametrize("junk", ["nonsense", 99, -1, object(), False])
def test_resolve_protection_junk_degrades_to_open(junk):
    # A malformed / unknown value must never raise — it degrades to OPEN (no mark).
    assert resolve_protection("n", {"protection": {"n": junk}}) is GuardLevel.OPEN


def test_is_protected_threshold():
    assert is_protected(GuardLevel.OPEN) is False
    assert is_protected(GuardLevel.GUARDED) is True
    assert is_protected(GuardLevel.PUBLISHED) is True


# ---------------------------------------------------------------------------
# LaTeX
# ---------------------------------------------------------------------------

def test_latex_empty_map_is_byte_identical(scaffold):
    baseline = generate_latex(scaffold)
    assert generate_latex({**scaffold, "protection": {}}) == baseline
    # Absent key already equals baseline by construction; assert explicitly.
    assert generate_latex(scaffold) == baseline


def test_latex_mark_appears_only_on_guarded_node(scaffold, guarded):
    out = generate_latex(guarded)
    assert f"@protected node={GUARDED_ID} level=GUARDED" in out
    assert f"@protected node={OTHER_ID}" not in out
    # The zero-arg wrapper commands + their providecommand identity are present
    # so the .tex compiles whether or not the class file overrides them.
    assert r"\rlmprotectedbegin" in out
    assert r"\rlmprotectedend" in out
    assert r"\providecommand{\rlmprotectedbegin}" in out


def test_latex_published_level_name_surfaced(scaffold):
    out = generate_latex({**scaffold, "protection": {GUARDED_ID: GuardLevel.PUBLISHED}})
    assert f"@protected node={GUARDED_ID} level=PUBLISHED" in out


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

def test_html_empty_map_is_byte_identical(scaffold):
    baseline = generate_html(scaffold)
    assert generate_html({**scaffold, "protection": {}}) == baseline


def test_html_mark_appears_only_on_guarded_node(scaffold, guarded):
    out = generate_html(guarded)
    assert 'class="protected protected-guarded"' in out
    assert 'data-protection-level="guarded"' in out
    # Exactly one protected wrapper (the single guarded node).
    assert out.count('class="protected ') == 1


def test_html_published_class(scaffold):
    out = generate_html({**scaffold, "protection": {GUARDED_ID: GuardLevel.PUBLISHED}})
    assert 'class="protected protected-published"' in out


# ---------------------------------------------------------------------------
# docx
# ---------------------------------------------------------------------------

def _shaded_paragraph_count(docx_bytes: bytes) -> int:
    doc = Document(BytesIO(docx_bytes))
    n = 0
    for p in doc.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:shd")) is not None:
            n += 1
    return n


def _docx_document_xml(docx_bytes: bytes) -> bytes:
    """The word/document.xml payload — the meaningful content, WITHOUT the zip
    envelope. A .docx is a zip and its entries embed wall-clock mod-times (2s
    granularity), so comparing raw archive bytes is timing-flaky; comparing the
    document part is both timestamp-free and what "byte-identical output" actually
    means for our purposes."""
    import zipfile

    with zipfile.ZipFile(BytesIO(docx_bytes)) as z:
        return z.read("word/document.xml")


def test_docx_empty_map_is_content_identical(scaffold):
    # Compare the document part, not the zip bytes (which carry flaky timestamps).
    baseline = _docx_document_xml(generate_docx(scaffold))
    with_empty = _docx_document_xml(generate_docx({**scaffold, "protection": {}}))
    assert with_empty == baseline


def test_docx_mark_appears_when_guarded(scaffold, guarded):
    baseline = generate_docx(scaffold)
    marked = generate_docx(guarded)
    # The document content differs (compared timestamp-free), the doc still
    # re-opens, and the guarded node's paragraphs are now shaded (none were in the
    # baseline).
    assert _docx_document_xml(marked) != _docx_document_xml(baseline)
    assert _shaded_paragraph_count(baseline) == 0
    assert _shaded_paragraph_count(marked) >= 1
    # Still a valid, re-openable .docx.
    assert Document(BytesIO(marked)) is not None


# ---------------------------------------------------------------------------
# JATS / BITS
# ---------------------------------------------------------------------------

def test_jats_empty_map_is_byte_identical(scaffold):
    baseline = generate_jats(scaffold)
    assert generate_jats({**scaffold, "protection": {}}) == baseline


def test_jats_mark_appears_only_on_guarded_node(scaffold, guarded):
    out = generate_jats(guarded)
    assert f"protected id={GUARDED_ID} level=guarded" in out
    assert f"protected id={OTHER_ID} " not in out


def test_bits_empty_map_is_byte_identical(scaffold):
    baseline = generate_bits(scaffold)
    assert generate_bits({**scaffold, "protection": {}}) == baseline


def test_bits_mark_appears_when_guarded(scaffold, guarded):
    out = generate_bits(guarded)
    assert f"protected id={GUARDED_ID} level=guarded" in out
