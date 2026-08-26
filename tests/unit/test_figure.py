"""
test_figure.py — the first-class `figure` semantic content type (ADR-0012).

Covers: the catalog entry + subtype vocabulary + graphic-role selection; the
positional figure_number pass; and that all three renderers embed the artifact
and render a "Figure N." caption from one shared node + payload shape.
"""

import base64
from io import BytesIO

import pytest
from docx import Document
from docx.oxml.ns import qn

import document_model.render_capabilities as rc
import rendering.docx_generator as dg
import rendering.latex_generator as lg
import rendering.html_generator as hg
from document_model.document_node import DocNode
from document_model.document_tree import compute_table_numbers, compute_figure_numbers


# A minimal VALID 4x4 PNG, base64 (no data-URI prefix) — must be a real PNG so
# python-docx's add_picture can read its dimensions.
_PNG_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAQAAAAECAIAAAAmkwkpAAAAEElEQVR4nGM8wYAATAxEcQAo3ADQ"
    "QD5s4QAAAABJRU5ErkJggg=="
)


def _figure_node(**kw) -> DocNode:
    defaults = dict(id="fig-1", title="Demo", node_type="figure",
                    subtype="chart", data_key="demo_fig",
                    caption="A demonstration figure.")
    defaults.update(kw)
    return DocNode(**defaults)


def _payload() -> dict:
    return {"png_b64": _PNG_B64, "filename": "demo.png",
            "caption": "A demonstration figure."}


# ---------------------------------------------------------------------------
# Catalog + subtype vocabulary
# ---------------------------------------------------------------------------

def test_figure_is_a_captionable_orientable_catalog_type():
    spec = rc.component_for("figure")
    assert rc.is_captionable("figure")
    assert spec.capabilities.orientable and spec.capabilities.breakable
    assert spec.headingless
    assert set(spec.content_kinds) == {"chart", "image"}


def test_figure_emits_furniture_roles():
    emits = rc.emits_for("figure")
    assert set(emits) == {"fig_title", "fig_caption", "fig_source",
                          "fig_note", "fig_alt_text"}


def test_figure_subtypes_and_graphic_role():
    assert rc.FIGURE_SUBTYPES == {"chart", "logo"}
    assert rc.figure_graphic_role("chart") == "fig_graphic"
    assert rc.figure_graphic_role("logo") == "logo_graphic"
    assert rc.figure_graphic_role(None) == "fig_graphic"


def test_figure_furniture_roles_resolve_in_the_vocabulary():
    import document_model.vocabulary as V
    vocab = V.load_vocabulary("ntp-report")
    for role in rc.emits_for("figure") + ("fig_graphic", "logo_graphic"):
        assert vocab.get(role) is not None, f"figure role {role!r} unresolved"


# ---------------------------------------------------------------------------
# Positional figure numbering
# ---------------------------------------------------------------------------

def test_figure_numbers_are_positional_and_separate_from_tables():
    tree = [_figure_node(id="f1"), _figure_node(id="f2")]
    compute_figure_numbers(tree)
    assert [n.figure_number for n in tree] == [1, 2]
    assert all(n.table_number is None for n in tree)


def test_compute_table_numbers_also_assigns_figure_numbers():
    # The figure pass is folded into compute_table_numbers so every caller gets both.
    fig = _figure_node()
    compute_table_numbers([fig])
    assert fig.figure_number == 1


# ---------------------------------------------------------------------------
# Rendering — all three surfaces, one node + payload shape
# ---------------------------------------------------------------------------

def test_html_figure_embeds_data_uri_and_numbered_caption():
    fig = _figure_node(); fig.figure_number = 1
    html = hg._render_figure(fig, {"demo_fig": _payload()})
    assert 'data:image/png;base64,' in html
    assert "Figure 1. A demonstration figure." in html
    # The alt text is the DESCRIPTIVE caption alone (accessibility role), no "Figure N.".
    assert 'alt="A demonstration figure."' in html


def test_latex_figure_includegraphics_and_numbered_caption():
    fig = _figure_node(); fig.figure_number = 2
    tex = lg._render_figure(fig, {"demo_fig": _payload()})
    assert r"\includegraphics" in tex and "figures/demo.png" in tex
    assert "Figure 2. A demonstration figure." in tex


def test_docx_figure_embeds_image_and_numbered_caption():
    fig = _figure_node(); fig.figure_number = 3
    doc = Document()
    dg._render_figure(doc, fig, {"demo_fig": _payload()})
    images = [r for r in doc.part.rels.values() if "image" in r.reltype]
    assert len(images) == 1
    assert any("Figure 3. A demonstration figure." in p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# Missing-payload → visible pending note, never a silent gap
# ---------------------------------------------------------------------------

def test_missing_payload_renders_pending_on_all_surfaces():
    fig = _figure_node()
    assert "pending" in hg._render_figure(fig, {}).lower()
    assert "pending" in lg._render_figure(fig, {}).lower()
    doc = Document(); dg._render_figure(doc, fig, {})
    assert any("pending" in p.text.lower() for p in doc.paragraphs)
