"""
test_inline_content.py — the minimal inline-content model (ext-link) and its
translation on all three surfaces (ADR-0010 inline sibling of the block catalog).

A hyperlink is SEMANTIC content (a display text + a target URI), not a styling
annotation; each surface derives its own presentation.  A paragraph is either a
plain str (unchanged) or a list of inline units (str runs + typed dicts).
"""

import rendering.render_common as rc
import rendering.html_generator as hg
import rendering.latex_generator as lg
import rendering.docx_generator as dg
from docx import Document
from docx.oxml.ns import qn


_PARA = [
    "Published by the ",
    rc.make_ext_link("NIEHS", "https://www.niehs.nih.gov/"),
    " and indexed in ",
    rc.make_ext_link("PubMed", "https://pubmed.ncbi.nlm.nih.gov/"),
    ".",
]


# ---------------------------------------------------------------------------
# The model in render_common
# ---------------------------------------------------------------------------

def test_normalize_inline_coerces_str_and_list():
    assert rc.normalize_inline("hello") == ["hello"]
    assert rc.normalize_inline("") == []
    assert rc.normalize_inline(None) == []
    assert rc.normalize_inline(_PARA) == _PARA


def test_inline_plain_text_flattens():
    assert rc.inline_plain_text("hello") == "hello"
    assert rc.inline_plain_text(_PARA) == "Published by the NIEHS and indexed in PubMed."


def test_paragraph_has_inline_only_for_typed_units():
    assert rc.paragraph_has_inline(_PARA) is True
    assert rc.paragraph_has_inline("plain") is False
    assert rc.paragraph_has_inline(["a", "b"]) is False   # plain runs, no typed unit


def test_has_paragraph_content_handles_inline_lists():
    assert rc.has_paragraph_content([_PARA]) is True
    assert rc.has_paragraph_content([[""]]) is False
    assert rc.has_paragraph_content(["plain"]) is True


# ---------------------------------------------------------------------------
# The three surface translations — a link is a link on every surface
# ---------------------------------------------------------------------------

def test_html_renders_anchor():
    out = hg._render_inline(_PARA)
    assert '<a href="https://www.niehs.nih.gov/">NIEHS</a>' in out
    assert '<a href="https://pubmed.ncbi.nlm.nih.gov/">PubMed</a>' in out


def test_latex_renders_href():
    out = lg._render_inline(_PARA)
    assert r"\href{https://www.niehs.nih.gov/}{NIEHS}" in out
    assert r"\href{https://pubmed.ncbi.nlm.nih.gov/}{PubMed}" in out


def test_docx_renders_hyperlink_relationships():
    doc = Document()
    para = doc.add_paragraph()
    dg._add_inline_runs(para, _PARA)
    links = para._p.findall(qn("w:hyperlink"))
    assert len(links) == 2
    targets = {r.target_ref for r in doc.part.rels.values() if "hyperlink" in r.reltype}
    assert "https://www.niehs.nih.gov/" in targets
    assert "https://pubmed.ncbi.nlm.nih.gov/" in targets
    # The display text is present, the URL is NOT inlined as text (it's the target).
    assert para.text == "Published by the NIEHS and indexed in PubMed."


def test_plain_string_paragraph_unchanged_on_all_surfaces():
    # A plain string keeps the exact pre-inline path (no <a>/\href/hyperlink).
    assert hg._render_inline("just prose") == "just prose"
    assert lg._render_inline("just prose") == "just prose"
    doc = Document(); p = doc.add_paragraph(); dg._add_inline_runs(p, "just prose")
    assert not p._p.findall(qn("w:hyperlink"))
    assert p.text == "just prose"


def test_unknown_inline_unit_degrades_to_text():
    para = ["see ", {"type": "xref", "text": "Table 1", "target": "t1"}]
    assert "Table 1" in hg._render_inline(para)
    assert "Table 1" in lg._render_inline(para)


# ---------------------------------------------------------------------------
# End-to-end: the foreword boilerplate carries its anchors on every surface
# ---------------------------------------------------------------------------

def test_foreword_hyperlinks_render_end_to_end():
    from rendering.report_data import scaffold_report_data
    from io import BytesIO
    data = scaffold_report_data(chemical_name="X", casrn="1-1-1", dtxsid="X")

    html = hg.generate_html(data)
    assert '<a href="https://www.niehs.nih.gov/">' in html
    assert 'href="https://pubmed.ncbi.nlm.nih.gov/"' in html

    tex = lg.generate_latex(data)
    assert r"\href{https://www.niehs.nih.gov/}" in tex

    doc = Document(BytesIO(dg.generate_docx(data)))
    targets = {r.target_ref for r in doc.part.rels.values() if "hyperlink" in r.reltype}
    assert {
        "https://www.niehs.nih.gov/",
        "https://www.niehs.nih.gov/research/atniehs/dtt/index.cfm",
        "https://pubmed.ncbi.nlm.nih.gov/",
    } <= targets
