"""Dev helper: build a STYLES-ONLY docx base from the NIEHS-10 reference.

Opens examples/NIEHS-10 PFHxSAm_Final.docx, strips ALL body content
(paragraphs + tables + everything except the final section's sectPr), and saves
assets/templates/niehs-10-base.docx.  python-docx carries styles.xml,
numbering.xml, theme1.xml, fontTable.xml, settings.xml forward automatically, so
the result is the reference's full 386-style library with an empty body — the
template base docx_generator opens instead of a blank Document().

Reproducible + reviewable: re-run to regenerate the base from the reference.
Not part of the shipped surface.
"""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

SRC = Path("examples/NIEHS-10 PFHxSAm_Final.docx")
OUT = Path("assets/templates/niehs-10-base.docx")


def strip_body(doc) -> None:
    """Remove every top-level body child EXCEPT the trailing sectPr, then leave a
    single empty paragraph so the body is well-formed."""
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    # Remove all children (paragraphs, tables, bookmarks, etc.).
    for child in list(body):
        body.remove(child)
    # One empty paragraph so the doc has a valid body; generate_docx clears it.
    from docx.oxml import OxmlElement
    body.append(OxmlElement("w:p"))
    # Restore the section properties last (page geometry / grid live here).
    if sectPr is not None:
        body.append(sectPr)


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"reference not found: {SRC}")
    doc = Document(str(SRC))
    n_styles = len(doc.styles)
    strip_body(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))

    # Verify the saved base: styles survived, body is empty.
    check = Document(str(OUT))
    body_paras = len(check.paragraphs)
    body_tables = len(check.tables)
    print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")
    print(f"  styles: {n_styles} -> {len(check.styles)}")
    print(f"  body: {body_paras} paragraphs, {body_tables} tables (expect ~1 / 0)")
    # Spot-check a few NTP named styles are present.
    for name in ("0-03_Paragraph", "3-02a_Head1_NoNumber", "0-25_Table_Title",
                 "toc 1", "table of figures"):
        try:
            check.styles[name]
            print(f"  OK style present: {name!r}")
        except KeyError:
            print(f"  MISSING style: {name!r}")


if __name__ == "__main__":
    main()
