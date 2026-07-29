r"""
docx_generator.py — Microsoft Word (.docx) rendering of the NIEHS biological
potency report.

This module is the third sibling of latex_generator.py and html_generator.py.
All three walk the same DOCUMENT_TREE with the same data dict and dispatch on the
same set of node_type values (render_common.RENDERABLE_NODE_TYPES); they differ
only in the output each per-node handler produces.  That gives a Word export a
1:1 semantic match with the Overleaf .tex and the in-app HTML preview — same
tree, same data, same dispatch, just a different emitter.

Why a third surface
-------------------
Some stakeholders' publishing workflows are Word-driven (track-changes review,
house .dotx templates).  ADR-0008 (Option B) accepted a Word emitter as a
one-way OUTPUT surface: we generate a .docx from the tree, but never treat Word
as a source of truth — the tree + data remain canonical (Architectural
Invariant #1/#2).  Round-tripping Word back INTO the tree is a separate,
existing concern (freeform_content.py ingests an authored .docx block); this
module is generation only.

Architecture mirror + the one deliberate divergence
---------------------------------------------------
The dispatch table and per-type handlers follow html_generator.py's structure.
The ONE structural difference: LaTeX and HTML emitters return markup STRINGS and
share render_common.walk_emit (which accumulates a flat list of chunks).  Word is
an object model — python-docx handlers MUTATE a Document in place — so a handler
here has signature ``_render_<type>(doc, node, data) -> None`` and this module
runs its own ``_walk_docx_tree`` rather than walk_emit.  Everything semantic
(which content source wins, table rows, captions, roster) still comes from the
shared render_common EXTRACT plans, so this surface cannot drift from the other
two on WHAT a section contains — only on how Word renders it.

Scope (v1)
----------
Full node-type parity (all 17 RENDERABLE_NODE_TYPES) + a clean typographic
cover/title page + roman→arabic two-section page numbering + booktabs-style
tables + inline genomics chart images.  NOT in v1 (deliberate, documented
follow-ups): the per-node ``layout_style`` styling the other two surfaces honour
(this surface uses the Word base styles), and a configurable-fonts binding.
"""

from __future__ import annotations

import base64
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from cover_layouts import get_cover_layout
from document_tree import DOCUMENT_TREE, DocNode, first_body_node_id, walk_tree
from freeform_content import pending_note as _freeform_pending_note
from render_common import (
    ANIMAL_ROSTER_HEADERS,
    BMD_SUMMARY_HEADERS,
    GENE_SET_TABLE_HEADERS,
    GENE_TABLE_HEADERS,
    apical_table_plan,
    appendix_roster_rows,
    assert_dispatch_covers,
    bmd_summary_plan,
    front_matter_plan,
    gene_set_table_rows,
    gene_table_rows,
    genomics_chart_caption,
    genomics_description_items,
    genomics_entries,
    genomics_intro_paragraphs,
    genomics_role,
    genomics_table_caption,
    has_paragraph_content,
    normalize_inline,
    inline_plain_text,
    INLINE_EXT_LINK,
    incidence_table_plan,
    methods_subsection_content,
    sample_counts_table,
    table_caption as _table_caption,
    unified_narrative_paragraphs,
)
from docx.opc.constants import RELATIONSHIP_TYPE as _REL
from genomics_content import genomics_content_plan
from layout_style import resolve_layout_style
from render_capabilities import front_matter_roles_for, landscape_requested

_REL_HYPERLINK = _REL.HYPERLINK
from table_builder_common import format_display_number, format_mean_se_display


# ---------------------------------------------------------------------------
# Constants — the reference typography spec
# ---------------------------------------------------------------------------
# Every value below was MEASURED from docs/NIEHS-Report-10-Reference.pdf (all 80
# pages, via PyMuPDF span analysis), so the generated Word styles reproduce the
# reference's fonts/sizes rather than Word's default theme (Calibri/Aptos).  We
# name the font FAMILIES ("Times New Roman", "Arial"); the actual glyphs come
# from the render machine's installed fonts — the reference-identical set lives
# in assets/fonts/ (gitignored, licensed) for local measurement only.

# US-Letter page + one-inch margins (NIEHS Report 10 trim: text block L=72pt,
# R≈540pt, top/bottom ≈35pt band for the running header/footer → 1" margins with
# the header/footer 0.5" from the edge).
_PAGE_WIDTH = Inches(8.5)
_PAGE_HEIGHT = Inches(11.0)
_MARGIN = Inches(1.0)
_HEADER_FOOTER_DISTANCE = Inches(0.5)

# Font families (matched by name against the render machine's installed fonts).
_BODY_FONT = "Times New Roman"     # body text — 78/80 reference pages
_HEADING_FONT = "Arial"            # all headings + table headers

# Point sizes (measured).  Body 12pt; headings step 17/15/13 by level; the
# front-matter 16/14pt heading variants collapse to level-1 17pt in v1 (a ≤1pt
# divergence that doesn't move pagination).  Tables start at a uniform 10pt (the
# reference steps 10→9→8.5→8 by density — a later diff-driven tuning pass).
_BODY_PT = 12
_HEADING_PT: dict[int, int] = {1: 17, 2: 15, 3: 13}
_TABLE_PT = 10
_HEADER_PT = 12

# Black — the reference body + title-page text color (headings and title).
_BLACK = RGBColor(0x00, 0x00, 0x00)

# DocNode.level → Word built-in heading style.  Level 0 = no heading (cover,
# leaf table nodes); levels 1-3 nest under the document title.
_HEADING_STYLE_BY_LEVEL: dict[int, str] = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
}

# Abstract font-family → concrete Word font, used only when a resolved style
# carries `font_family` (serif/sans/mono) but NOT an explicit `font` name.  The
# literal `font` always wins (see layout_style.LAYOUT_KEY_SCHEMA precedence note).
_DOCX_FONT_BY_FAMILY: dict[str, str] = {
    "serif": "Times New Roman",
    "sans": "Arial",
    "mono": "Consolas",
}

# Points per absolute unit, for converting a resolved-style length to Pt.  em/ex
# are relative and can't be resolved without the current font size, so a length
# in those units is ignored on the docx surface (logged as skipped).
_PT_PER_UNIT: dict[str, float] = {
    "pt": 1.0, "mm": 72.0 / 25.4, "cm": 72.0 / 2.54, "in": 72.0,
}
_LENGTH_RE = re.compile(r"^(-?\d+(?:\.\d+)?)(pt|mm|cm|in|em|ex)$")

# Strip C0 control chars (except tab/newline) that would make the OOXML invalid.
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
# Minimal HTML tag stripper for freeform html fallback (see _freeform_text).
_TAG_RE = re.compile(r"<[^>]+>")


def _clean(text) -> str:
    """Coerce a value to a control-char-free string safe for an OOXML run."""
    return _CONTROL_RE.sub("", str(text if text is not None else ""))


def _style_run(run, *, size_pt: int, font: str | None = None, bold: bool = False):
    """Set explicit size/font/weight on a run (used for table cells, which must
    not inherit the 12pt body size — the reference tables run smaller)."""
    run.font.size = Pt(size_pt)
    if font:
        run.font.name = font
    if bold:
        run.bold = True
    return run


# ---------------------------------------------------------------------------
# Low-level docx helpers
# ---------------------------------------------------------------------------

def _add_heading(doc: Document, level: int, title: str, data: dict | None = None):
    """Append a heading paragraph for levels 1-3; level 0/empty adds nothing.

    Returns the created paragraph (or None when nothing was added) so a caller can
    post-adjust it — e.g. front-matter/tables-list headings center theirs.

    Role-driven path (ADR-0010 Phase 2): when a vocabulary is active (``data``
    carries it), the level maps to the section_heading_N role and the paragraph
    gets that role's NATIVE Word style (3-0Na_HeadN_NoNumber).  Otherwise the
    built-in Heading 1-3 styles (_HEADING_STYLE_BY_LEVEL) — the pre-vocabulary
    look."""
    if not title or level not in _HEADING_STYLE_BY_LEVEL:
        return None
    style = _HEADING_STYLE_BY_LEVEL[level]
    if data is not None:
        role_style = _role_style_name(data, f"section_heading_{level}")
        if role_style and role_style in {s.name for s in doc.styles}:
            style = role_style
    return doc.add_paragraph(_clean(title), style=style)


def _add_paragraphs(doc: Document, paragraphs, style: str | None = None) -> bool:
    """
    Append one body paragraph per non-empty string.  Returns True if anything
    was added (so callers can fall back to a pending line when nothing was).

    ``style`` names a Word paragraph style to apply (role-driven path, e.g.
    1-22_Foreword_Text); None uses the default Normal (legacy path).  A style
    name absent from the doc is ignored (falls back to Normal), so an unbuilt
    vocabulary never breaks generation.
    """
    added = False
    use_style = style if (style and style in {s.name for s in doc.styles}) else None
    for p in paragraphs or []:
        # A paragraph is a plain str OR a list of inline units (render_common
        # inline model).  Skip only when it has no text at all.
        if not inline_plain_text(p).strip():
            continue
        para = doc.add_paragraph()
        if use_style:
            para.style = doc.styles[use_style]
        _add_inline_runs(para, p)
        added = True
    return added


def _add_inline_runs(paragraph, content) -> None:
    """Populate a paragraph from an inline-content value (a plain str OR a list of
    inline units — render_common's model).  A plain str is one run; an ext-link
    unit becomes a real Word hyperlink (blue, underlined) via _add_hyperlink; an
    unknown typed unit degrades to a plain run of its text."""
    for unit in normalize_inline(content):
        if isinstance(unit, str):
            if unit:
                paragraph.add_run(_clean(unit))
        elif unit.get("type") == INLINE_EXT_LINK:
            _add_hyperlink(paragraph, _clean(unit.get("text", "")), unit.get("href", ""))
        else:
            paragraph.add_run(_clean(unit.get("text", "")))


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Append a clickable external hyperlink run to a paragraph.

    python-docx has no hyperlink API, so build the OOXML directly: register the
    URL as an EXTERNAL relationship on the part, then a <w:hyperlink r:id=...>
    wrapping a run styled blue + underlined (the reference's link presentation).
    A missing url degrades to a plain run (the link text, not clickable)."""
    if not url:
        paragraph.add_run(text)
        return
    r_id = paragraph.part.relate_to(
        url, _REL_HYPERLINK, is_external=True
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rpr.append(color)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _add_labeled_sections(doc: Document, parts) -> None:
    """
    Structured-abstract parts as paragraphs with a bold run-in label — the Word
    twin of html_generator._render_labeled_sections.  A "" label is unlabeled.
    """
    for label, text in parts:
        para = doc.add_paragraph()
        if label:
            run = para.add_run(f"{_clean(label)}. ")
            run.bold = True
        para.add_run(_clean(text))


def _add_pending(doc: Document, label: str) -> None:
    """Visible italic "[pending]" placeholder — the Word twin of _pending."""
    para = doc.add_paragraph()
    run = para.add_run(f"[{_clean(label)}]")
    run.italic = True


def _set_repeat_header(row) -> None:
    """Mark a table row as a header that repeats on each page (tblHeader)."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _set_edge_border(el, edge: str, sz: str = "8") -> None:
    """
    Give a table (tblPr/tblBorders) or cell (tcPr/tcBorders) a single ``edge``
    border (top/bottom/...), leaving the others untouched.  Used to build the
    booktabs look (horizontal rules only).
    """
    if el.tag.endswith("}tbl"):
        pr = el.tblPr
        borders_tag = "w:tblBorders"
    else:  # a <w:tc> cell
        pr = el.get_or_add_tcPr()
        borders_tag = "w:tcBorders"
    borders = pr.find(qn(borders_tag))
    if borders is None:
        borders = OxmlElement(borders_tag)
        pr.append(borders)
    e = borders.find(qn(f"w:{edge}"))
    if e is None:
        e = OxmlElement(f"w:{edge}")
        borders.append(e)
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), sz)
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), "000000")


# The reference table-caption style + Table-of-Figures collect target.  A caption
# in this style, with a SEQ number, is what the front-matter "Tables" TOF field
# gathers (see _render_tables_list).  Values measured from NIEHS-10's
# 0-25_Table_Title (Base_Text → 11pt bold, 14pt before / 3pt after, keepNext).
_TABLE_TITLE_STYLE = "0-25_Table_Title"
_TABLE_TITLE_PT = 11


def _ensure_table_title_style(doc: Document):
    """Create (once) the `0-25_Table_Title` paragraph style — the reference table-
    caption style AND the target the Tables-list TOF field collects by.  Times New
    Roman 11pt bold, 14pt before / 3pt after, keepNext (stay with the table)."""
    name = _TABLE_TITLE_STYLE
    try:
        return doc.styles[name]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = _BODY_FONT
        style.font.size = Pt(_TABLE_TITLE_PT)
        style.font.bold = True
        pf = style.paragraph_format
        pf.space_before = Pt(14)
        pf.space_after = Pt(3)
        pf.keep_with_next = True
        return style


def _add_table_caption(doc: Document, caption: str) -> None:
    r"""Emit a table caption ABOVE its table in `0-25_Table_Title` style, numbered
    by a SEQ FIELD ("Table {SEQ Table}. <text>") rather than literal text — so the
    number auto-renumbers AND the front-matter Table-of-Figures field can collect
    it (a TOF collects by paragraph style).  The incoming caption already carries a
    literal "Table N. " prefix (from render_common.table_caption); strip it and let
    the SEQ supply the number.  A non-breaking space after "Table" keeps the label
    with its number (matching the reference)."""
    style = _ensure_table_title_style(doc)
    para = doc.add_paragraph(style=style)
    label, text = _split_table_prefix(_clean(caption))
    if label and _re_arabic_label.match(label):
        # Body table: "Table N." -> SEQ field cached to N (correct positional
        # number so LibreOffice shows it right without a field refresh).
        para.add_run("Table ")   # nbsp, so "Table" never wraps off its number
        _add_seq_field(para, "Table", cached=label)
        para.add_run(". ")
        para.add_run(text)
    elif label:
        # Appendix table: "Table B-1." -> literal (letter+number the arabic SEQ
        # can't express); emitting a SEQ here is what doubled the prefix.
        para.add_run(f"Table {label}. ")
        para.add_run(text)
    else:
        # No recognizable prefix -> plain SEQ-numbered caption.
        para.add_run("Table ")
        _add_seq_field(para, "Table")
        para.add_run(". ")
        para.add_run(text)


# "N"  (body, arabic) vs  "B-1" (appendix, letter-prefixed).
_re_arabic_label = re.compile(r"^\d+$")


def _split_table_prefix(caption: str) -> tuple[str, str]:
    """Split a caption into (number-label, descriptive-text).

    Recognizes both the body form ("Table 3. Foo" -> ("3", "Foo")) and the
    appendix form ("Table B-1. Bar" -> ("B-1", "Bar")).  Returns ("", caption)
    when there is no recognizable "Table .... " prefix.  The SEQ field / TOF
    numbering then supplies the label from the returned number for body tables;
    appendix labels are emitted literally (see _add_table_caption)."""
    m = re.match(r"^Table\s+([A-Za-z]?-?\d+)\.\s*(.*)$", caption or "", re.DOTALL)
    if m:
        return m.group(1), m.group(2).strip()
    return "", (caption or "").strip()


def _strip_table_prefix(caption: str) -> str:
    """Drop a leading 'Table N. ' (or 'Table B-1. ') from a caption — the SEQ
    field / TOF numbering supplies the label, so the stored text is just the
    descriptive part.  Thin wrapper over _split_table_prefix for callers that
    only want the text."""
    return _split_table_prefix(caption)[1]


def _add_seq_field(paragraph, seq_name: str, cached: str = "1") -> None:
    r"""Insert a Word SEQ field ({ SEQ <name> \* ARABIC }) — an auto-incrementing
    counter Word renumbers on update.  `cached` is the result shown until a field
    refresh; set it to the correct positional number so LibreOffice (which never
    refreshes a Word-origin field) still shows the right label."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {seq_name} \\* ARABIC "
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    cached_t = OxmlElement("w:t"); cached_t.text = cached
    cached_run = OxmlElement("w:r"); cached_run.append(cached_t)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(sep)
    paragraph._p.append(cached_run)
    tail = paragraph.add_run(); tail._r.append(end)


def _booktabs_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    caption: str = "",
    footnotes: list | None = None,
    merged_label_rows: bool = False,
) -> None:
    """
    Append a booktabs-style table (top rule, header rule, bottom rule; no
    vertical or inner-body rules) — the Word analogue of the ``niehstable``
    environment / HTML ``.niehstable``.

    caption:            rendered ABOVE the table (NIEHS places captions above) in
                        the reference `0-25_Table_Title` style with a SEQ-field
                        number, so it is COLLECTABLE by the front-matter Table-of-
                        Figures field and auto-renumbers (see _add_table_caption).
    footnotes:          the typed footnote dicts render as small paragraphs
                        below the table (lettered entries keep their letter).
    merged_label_rows:  when True, a body row whose first cell is wrapped in
                        ``**...**`` (the apical / sample-counts sex separator)
                        is emitted as a single bold cell merged across all
                        columns — the Word twin of the HTML colspan row.
    """
    if caption:
        _add_table_caption(doc, caption)

    ncols = max(len(headers), max((len(r) for r in rows), default=0), 1)
    table = doc.add_table(rows=0, cols=ncols)
    table.autofit = True

    if headers:
        hcells = table.add_row().cells
        for i, h in enumerate(headers):
            if i >= ncols:
                break
            para = hcells[i].paragraphs[0]
            # Reference table headers are Arial Bold, one step below body size.
            _style_run(para.add_run(_clean(h)), size_pt=_TABLE_PT,
                       font=_HEADING_FONT, bold=True)
        _set_repeat_header(table.rows[0])
        for c in table.rows[0]._tr.tc_lst:
            _set_edge_border(c, "bottom")

    for row in rows:
        cells = [_clean(c) for c in row]
        first = cells[0] if cells else ""
        if merged_label_rows and first.startswith("**") and first.endswith("**"):
            tr = table.add_row()
            merged = tr.cells[0]
            for other in tr.cells[1:]:
                merged = merged.merge(other)
            _style_run(merged.paragraphs[0].add_run(first.strip("*").strip()),
                       size_pt=_TABLE_PT, bold=True)
            continue
        tr = table.add_row()
        for i in range(ncols):
            val = cells[i].strip() if i < len(cells) else ""
            # Data cells: body font (Times), table size (10pt).
            _style_run(tr.cells[i].paragraphs[0].add_run(val), size_pt=_TABLE_PT)

    # Outer top + bottom rules complete the booktabs frame.
    _set_edge_border(table._tbl, "top")
    _set_edge_border(table._tbl, "bottom")

    _add_footnotes(doc, footnotes)


def _add_footnotes(doc: Document, footnotes) -> None:
    """Render table footnotes as small paragraphs below a table."""
    for fn in footnotes or []:
        if not isinstance(fn, dict):
            continue
        text = fn.get("text") or fn.get("body") or ""
        if not text:
            continue
        para = doc.add_paragraph()
        if fn.get("kind") == "lettered" and fn.get("letter"):
            marker = para.add_run(f"{_clean(fn['letter'])} ")
            marker.font.superscript = True
            marker.font.size = Pt(8)
        run = para.add_run(_clean(text))
        run.font.size = Pt(8)


def _add_page_break(doc: Document) -> None:
    """Append an explicit page break (its own empty paragraph)."""
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Per-node-type render handlers — _render_<type>(doc, node, data) -> None
# ---------------------------------------------------------------------------
# Each handler appends this node's OWN content to the Document (children are
# walked separately by _walk_docx_tree).  Mirrors the html_generator handlers
# 1:1 in which shared EXTRACT plan it consumes.


def _render_front_matter(doc: Document, node: DocNode, data: dict) -> None:
    """Front-matter / narrative section: heading + labeled-sections or prose.

    Role-driven path (ADR-0010): a ``front-matter`` node derives its heading + body
    ROLES from its data_key (render_capabilities.front_matter_roles_for) — so the
    Foreword title is 1-21_Foreword_Title (centered), its text 1-22_Foreword_Text,
    the Abstract head 1-15_Abstract_Head, etc., instead of the generic body
    section_heading/Normal.  `narrative` nodes delegating here keep the generic
    heading (they are body sections, not front matter)."""
    heading_role = body_role = None
    if data.get("_vocabulary") is not None:
        if node.node_type == "front-matter":
            heading_role, body_role = front_matter_roles_for(node.data_key)
        else:
            # `narrative` body sections (Background / Summary / References …) keep
            # the generic section heading, but their prose is the canonical NTP
            # body paragraph style (0-03_Paragraph), not Normal.
            body_role = "body_para"

    heading_style = _role_style_name(data, heading_role) if heading_role else None
    if heading_style and node.title and heading_style in {s.name for s in doc.styles}:
        head = doc.add_paragraph(_clean(node.title), style=doc.styles[heading_style])
    else:
        head = _add_heading(doc, node.level, node.title, data)
    # Front-matter section headers get the reference front-matter heading look
    # (Arial 16pt bold, centered, 12pt after — matching the TOC/Tables headers).
    # Body sections (Background/Summary/References are `narrative` nodes that ALSO
    # route through here) stay as the left-aligned built-in Heading 1, matching
    # the NTP reference's left body 3-0Na heads — so gate on the node actually
    # being front-matter, not merely on the render path.  A role style (1-NN)
    # already carries its own typography, so only the built-in path needs this.
    if head is not None and not heading_style and node.node_type == "front-matter":
        _style_front_matter_heading(head)

    body_style = _role_style_name(data, body_role) if body_role else None
    plan = front_matter_plan(node, data)
    if plan.kind == "labeled":
        _add_labeled_sections(doc, plan.labeled_parts)
    elif plan.kind == "paragraphs":
        _add_paragraphs(doc, plan.paragraphs, style=body_style)
    else:
        _add_pending(doc, f"Section pending: {node.title}")


def _render_narrative(doc: Document, node: DocNode, data: dict) -> None:
    """Plain narrative — M&M subsections route through the methods lookup."""
    if node.methods_key:
        _render_methods_subsection(doc, node, data)
    else:
        _render_front_matter(doc, node, data)


def _render_methods_subsection(doc: Document, node: DocNode, data: dict) -> None:
    """M&M subsection — prose + optional inline table, matched by methods_key."""
    _add_heading(doc, node.level, node.title, data)
    paragraphs, inline = methods_subsection_content(node, data)
    body_style = _pstyle_or_default(doc, data, "body_para")
    added = (_add_paragraphs(doc, paragraphs, style=body_style)
             if has_paragraph_content(paragraphs) else False)
    if inline is not None:
        _render_inline_table(doc, inline)
        added = True
    if not added:
        _add_pending(doc, f"Section pending: {node.title}")


def _render_inline_table(doc: Document, table: dict) -> None:
    """Inline neutral {caption, headers, rows, footnotes} table."""
    headers = [str(h) for h in table.get("headers", [])]
    rows = [[str(c) for c in r] for r in table.get("rows", [])]
    if not headers and not rows:
        return
    _booktabs_table(
        doc, headers, rows,
        caption=table.get("caption", ""),
        footnotes=table.get("footnotes", []),
    )


def _render_sample_counts_table(doc: Document, node: DocNode, data: dict) -> None:
    """Methods 'Final Sample Counts' matrix (Table 1) with sex-separator rows."""
    built = sample_counts_table(node, data)
    if built is None:
        _add_pending(doc, f"Table data pending: {node.title}")
        return
    headers = [str(h) for h in built.get("headers", [])]
    rows = [[str(c) for c in r] for r in built.get("rows", [])]
    _booktabs_table(
        doc, headers, rows,
        caption=_table_caption(node, built.get("caption", "")),
        footnotes=built.get("footnotes", []),
        merged_label_rows=True,
    )


def _render_heading_only(doc: Document, node: DocNode, data: dict) -> None:
    """Structural heading; children rendered separately by the walker."""
    _add_heading(doc, node.level, node.title, data)


def _render_appendix(doc: Document, node: DocNode, data: dict) -> None:
    """Appendix — B renders the animal roster; others heading + stub/children."""
    _add_heading(doc, node.level, node.title, data)
    rows = appendix_roster_rows(node, data)
    if rows is not None:
        _booktabs_table(
            doc, list(ANIMAL_ROSTER_HEADERS), rows,
            caption="Table B-1. Animal Numbers and FASTQ Data File Names",
        )
        return
    if not node.children:
        _add_pending(doc, f"Appendix body pending: {node.title}")


# The Table-of-Figures instruction: collect paragraphs styled 0-25_Table_Title
# (the body table captions), as hyperlinks (\h), leaders hidden in web view (\z),
# caption/SEQ-based list (\c).  Byte-matches NIEHS-10's "Tables" field.
_TOF_INSTR = r'TOC \h \z \t "0-25_Table_Title" \c'
# The reference "table of figures" entry style: right dot-leader tab at the text
# edge (9360 twips = 6.5"), basedOn Normal.  Line/after measured from NIEHS-10's
# rendered PDF (see _ensure_tof_entry_style): tight 15.9pt within-entry line, 9.9pt
# after between entries → 41.7pt two-line / 25.8pt single-line entry pitch.
_TOF_ENTRY_STYLE = "table of figures"
_TOF_RIGHT_TAB_PT = 468


def _ensure_tof_entry_style(doc: Document):
    """Create (once) the `table of figures` entry style — mirroring the reference
    `table of figures` STYLE (examples/NIEHS-10 styles.xml): basedOn Normal, with
    a right dot-leader tab at the text-block edge, and NO spacing override of its
    own (line pitch comes from the section docGrid, ~18pt, exactly as the
    reference's does).  An earlier build FORCED a tight 15.9pt exact line + 9.9pt
    after to reverse-engineer the rendered pitch, but that diverged from the
    reference's actual grid-driven style, so it's removed."""
    try:
        return doc.styles[_TOF_ENTRY_STYLE]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        style = doc.styles.add_style(_TOF_ENTRY_STYLE, WD_STYLE_TYPE.PARAGRAPH)
        # basedOn Normal (the reference's model) so it inherits Normal's font +
        # grid-driven spacing rather than carrying its own line override.
        try:
            style.base_style = doc.styles["Normal"]
        except KeyError:
            pass
        style.font.name = _BODY_FONT
        style.font.size = Pt(_BODY_PT)
        spf = style.paragraph_format
        spf.tab_stops.add_tab_stop(
            Pt(_TOF_RIGHT_TAB_PT), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS,
        )
        return style


def _render_tables_list(doc: Document, node: DocNode, data: dict) -> None:
    """Front-matter list of tables — a native Word Table-of-Figures FIELD, the
    SAME field class as the Table of Contents (the reviewer's observation, borne
    out by NIEHS-10: its "Tables" section is `TOC \\h \\z \\t "0-25_Table_Title"
    \\c`, not a plain list).  The field COLLECTS every body caption styled
    0-25_Table_Title (see _add_table_caption) and renders "Table N.  <caption> …
    <page>".

    Pre-populated with cached entries for the same reason the TOC is (LibreOffice
    won't auto-refresh a Word-origin field); _mark_fields_dirty flips updateFields
    so Word offers to rebuild it, at which point it re-collects the real captions
    with real page numbers."""
    head = _add_heading(doc, node.level, node.title, data)
    _style_front_matter_heading(head)  # reference front-matter heading look
    entries = data.get("table_entries") or []
    if not entries:
        _add_pending(doc, "List of tables: pending.")
        return
    _add_tof_field(doc, entries)


def _add_tof_field(doc: Document, entries: list) -> None:
    r"""Append a pre-populated Table-of-Figures FIELD (see _add_toc_field for the
    same multi-paragraph field shape).  Instruction _TOF_INSTR; each cached entry
    is a `table of figures` paragraph "Table N.  <title>  \t —" (page-number
    placeholder, refreshed on update).  A well-formed empty field when there are
    no entries."""
    if not entries:
        para = doc.add_paragraph()
        r = para.add_run()
        for kind in ("begin", "separate", "end"):
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), kind)
            if kind == "begin":
                r._r.append(fc)
                instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
                instr.text = _TOF_INSTR
                r._r.append(instr)
            else:
                r._r.append(fc)
        return

    style = _ensure_tof_entry_style(doc)
    for i, entry in enumerate(entries):
        num = entry.get("table_number")
        title = _clean(entry.get("title", ""))
        para = doc.add_paragraph(style=style)

        if i == 0:
            opener = para.add_run()
            begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
            instr.text = _TOF_INSTR
            sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
            opener._r.append(begin); opener._r.append(instr); opener._r.append(sep)

        label = f"Table {num}. " if num is not None else ""
        para.add_run(f"{label}{title}")
        para.add_run("\t")
        para.add_run("—")  # page-number placeholder → real page on refresh

    _close_field_with_spacer(doc)


def _render_toc(doc: Document, node: DocNode, data: dict) -> None:
    """
    Contents — a NATIVE Word TOC field ({ TOC \\o "1-3" \\h \\z \\u }),
    PRE-POPULATED with cached entry paragraphs.

    A field is the medium-native answer, mirroring the LaTeX \\tableofcontents:
    Word/LibreOffice recompute it against the paginated document, so a reader who
    updates the field gets exact page numbers, dot leaders, and clickable links —
    which a static list can never carry (a page number does not exist until
    layout).  Collection is by OUTLINE LEVEL, so every body heading style carries
    outlineLvl 0-2 (built-in Heading 1-3 do; _build_vocabulary_styles stamps the
    NTP role styles to match).

    Why pre-populate: LibreOffice does NOT reliably honor <w:updateFields> for a
    Word-origin TOC field — it shows the field's CACHED content until the user
    runs Update Index (a known Word↔LibreOffice quirk).  An empty field therefore
    reads as a bare placeholder on open.  So we write the field's cached result
    the same way Word does when it saves: one TOC-N paragraph per heading between
    the field's `separate` and `end` marks.  The document opens showing a real
    contents list; updating the field then refreshes the page numbers in place.
    We still set <w:updateFields> (see _mark_fields_dirty) so Word offers to
    refresh; the cached content is what guarantees a non-empty TOC either way.

    The heading is a front-matter heading (restyled to the shared reference look:
    Arial 16pt bold, centered, 12pt after — see _style_front_matter_heading).  We
    resolve its base paragraph style through _pstyle_or_default so it uses a style
    that EXISTS in the doc (the NTP `1-23_FrontMatter_Head1` on the vocab/base
    path; plain Normal otherwise) — the old hardcoded "TOC Heading" name is a Word
    built-in absent from the NTP style base.
    """
    style_name = _pstyle_or_default(doc, data, "frontmatter_head1")
    head = (doc.add_paragraph(_clean(node.title), style=style_name)
            if style_name else doc.add_paragraph(_clean(node.title)))
    _style_front_matter_heading(head)
    entries = data.get("toc_entries") or []
    _add_toc_field(doc, entries)


# Right tab stop for a TOC line: the text width (US-Letter, 1" margins = 6.5").
# The page-number placeholder rides a dot-leader tab to this position, so the
# cached TOC mirrors Word's own dotted-leader layout.
_TOC_RIGHT_TAB_PT = 468

# The reference front-matter HEADING look, shared by every front-matter section
# header (Foreword, Table of Contents, Tables, About, Peer Review, Abstract, …).
# Measured from NIEHS-10's 1-23_FrontMatter_Head1 / NTP Contents Heading, which
# are identical: Arial 16pt bold, CENTERED, 12pt after.  Applying it uniformly
# fixes the drift where only the Contents header got the treatment and the others
# fell to built-in Heading 1 (17pt, 4pt-after → looked bigger AND jammed).
_FRONT_HEADING_PT = 16
_FRONT_HEADING_AFTER_PT = 12

# The reference `toc 1` style adds 6pt (spacing before=120 twips) before each
# top-level entry to group the section blocks; `toc 2`/`toc 3` carry no spacing.
# Line pitch itself comes from the section's docGrid (18pt), not a line override.
_TOC1_BEFORE_PT = 6


def _style_front_matter_heading(para) -> None:
    """Apply the reference front-matter heading look to a heading paragraph:
    Arial 16pt bold, centered, 12pt after (1-23_FrontMatter_Head1)."""
    if para is None:
        return
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(_FRONT_HEADING_AFTER_PT)
    for r in para.runs:
        r.font.name = _HEADING_FONT
        r.font.size = Pt(_FRONT_HEADING_PT)
        r.font.bold = True
        r.font.color.rgb = _BLACK


def _close_field_with_spacer(doc: Document) -> None:
    """Close a pre-populated TOC/TOF field the way the reference does: an EMPTY
    spacer paragraph whose paragraph-mark sits INSIDE the field (between separate
    and end), then the `end` fldChar on a following run.

    Why two steps rather than one closer paragraph: a reader (LibreOffice) shades
    the field region only up to the `end` MARK.  If `end` is the sole content of a
    paragraph, that paragraph's blank line is past the shaded region — the gray
    stops at the last entry (the reported bug) and the lone end-paragraph adds an
    extra blank line (the reported extra whitespace).  Putting the spacer's mark
    BEFORE `end` makes that blank line fall inside the field → it shades, matching
    the reference's one shaded blank line after the last entry; the `end` then
    rides the same paragraph so no second blank line is added."""
    spacer = doc.add_paragraph()
    # Empty run first (the shaded blank), then the end fldChar in a trailing run —
    # the paragraph-mark (and its blank line) is inside separate..end, so shaded.
    spacer.add_run("")
    end_run = spacer.add_run()
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _add_toc_field(doc: Document, entries: list) -> None:
    r"""Append a pre-populated Table-of-Contents FIELD spanning several paragraphs.

    python-docx has no field API, so build the OOXML directly.  The field
    instruction ``TOC \o "1-3" \h \z \u`` = collect outline levels 1-3 (\o), as
    hyperlinks (\h), suppress leaders in web layout (\z), use applied paragraph
    outline levels (\u).  Layout of the field across paragraphs (Word's own shape
    for a saved TOC):

        ¶  [fldChar begin][instrText TOC …][fldChar separate]  <entry 1 text>
        ¶  <entry 2 text>
        …
        ¶  <entry N text>[fldChar end]

    Each entry paragraph is styled TOC 1/2/3 by level and carries the title, a
    dot-leader tab, and a page-number PLACEHOLDER ("—"); the real number lands
    when the field is refreshed.  With no entries we still emit a well-formed
    (empty) field so a scaffold render doesn't crash."""
    if not entries:
        para = doc.add_paragraph()
        r = para.add_run()
        for kind in ("begin", "separate", "end"):
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), kind)
            if kind == "begin":
                r._r.append(fc)
                instr = OxmlElement("w:instrText")
                instr.set(qn("xml:space"), "preserve")
                instr.text = r'TOC \o "1-3" \h \z \u'
                r._r.append(instr)
            else:
                r._r.append(fc)
        return

    for i, entry in enumerate(entries):
        level = entry.get("level", 1)
        if not isinstance(level, int) or level < 1:
            level = 1
        level = min(level, 3)
        para = _toc_entry_paragraph(doc, level)

        if i == 0:
            # First entry paragraph opens the field (begin + instr + separate).
            opener = para.add_run()
            begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
            instr.text = r'TOC \o "1-3" \h \z \u'
            sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
            opener._r.append(begin)
            opener._r.append(instr)
            opener._r.append(sep)

        para.add_run(_clean(entry.get("title", "")))
        # Tab to the right margin (dot leader) + a page-number placeholder.
        para.add_run("\t")
        para.add_run("—")  # em dash — refreshed to the page number on update

    _close_field_with_spacer(doc)


def _toc_entry_paragraph(doc: Document, level: int):
    """A TOC entry paragraph styled TOC 1/2/3 (created on first use), with a
    right dot-leader tab at the text-block edge and level indentation.

    Spacing matches the reference `toc N` styles (examples/NIEHS-*.docx): DEFAULT
    single line spacing (no line-spacing override, no forced space_after — the
    earlier 0pt-after made entries tighter than the reference), with 6pt `before`
    on top-level (`toc 1`) entries only, which visually groups each section block.
    Baked onto the STYLE (the reference's model), so every entry inherits it.  The
    built-in TOC styles are absent from a fresh Document(), so add them lazily."""
    style_name = f"TOC {level}"
    try:
        style = doc.styles[style_name]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = _BODY_FONT
        style.font.size = Pt(_BODY_PT)
        spf = style.paragraph_format
        # Mirror the reference `toc N` STYLE definitions exactly (examples/
        # NIEHS-10 styles.xml): NO line-spacing override — every entry snaps to
        # the section's docGrid (linePitch 360 = 18pt), which is what gives the
        # reference its even line rhythm.  `toc 1` adds 6pt `before` (spacing
        # before=120) to group each top-level block; `toc 2`/`toc 3` carry NO
        # spacing at all.  An earlier build FORCED an exact 25.9pt line here to
        # reverse-engineer the rendered pitch — but that diverged from the
        # reference's actual style (grid-driven), so it's removed.
        if level == 1:
            spf.space_before = Pt(_TOC1_BEFORE_PT)
        if level > 1:
            spf.left_indent = Pt((level - 1) * 18)
    para = doc.add_paragraph(style=style)
    # Right-aligned tab with a dotted leader → the classic TOC dot leader.
    para.paragraph_format.tab_stops.add_tab_stop(
        Pt(_TOC_RIGHT_TAB_PT), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS,
    )
    return para


def _render_narrative_tables(doc: Document, node: DocNode, data: dict) -> None:
    """H2 group under Results: heading + unified narrative; tables walked after."""
    _add_heading(doc, node.level, node.title, data)
    paragraphs = unified_narrative_paragraphs(node, data)
    if has_paragraph_content(paragraphs):
        _add_paragraphs(doc, paragraphs, style=_pstyle_or_default(doc, data, "body_para"))
    else:
        _add_pending(doc, f"Narrative pending: {node.title}")


def _render_apical_table(doc: Document, node: DocNode, data: dict) -> None:
    """Apical dose-response table (Male/Female blocks, BMD/BMDL columns)."""
    plan = apical_table_plan(node, data)
    if plan is None:
        _add_pending(doc, f"Table data pending: {node.title}")
        return
    headers = (
        [plan.first_col]
        + [_format_dose_label(d, plan.dose_unit) for d in plan.doses]
        + [f"BMD 1Std ({plan.dose_unit})", f"BMDL 1Std ({plan.dose_unit})"]
    )
    rows: list[list[str]] = []
    for block in plan.sex_blocks:
        rows.append([f"**{block.sex_label}**"])
        for row in block.rows:
            rows.append(row.cells)
    _booktabs_table(
        doc, headers, rows,
        caption=plan.caption, footnotes=plan.footnotes, merged_label_rows=True,
    )


def _render_incidence_table(doc: Document, node: DocNode, data: dict) -> None:
    """Clinical Observations incidence table (observation × dose group)."""
    plan = incidence_table_plan(node, data)
    if plan is None:
        _add_pending(doc, f"Table data pending: {node.title}")
        return
    headers = ["Observation"] + [_format_dose_label(d, plan.dose_unit) for d in plan.doses]
    _booktabs_table(
        doc, headers, plan.rows, caption=plan.caption, footnotes=plan.footnotes,
    )


def _render_bmd_summary(doc: Document, node: DocNode, data: dict) -> None:
    """Apical Endpoint BMD Summary — prose + one row per endpoint."""
    plan = bmd_summary_plan(node, data)
    _add_heading(doc, node.level, node.title, data)
    prose = _add_paragraphs(doc, plan.paragraphs, style=_pstyle_or_default(doc, data, "body_para"))
    if plan.rows is None:
        if not prose:
            _add_pending(doc, f"BMD summary endpoints pending: {node.title}")
        return
    _booktabs_table(doc, list(BMD_SUMMARY_HEADERS), plan.rows, caption=plan.caption)


def _render_genomics_section(doc: Document, node: DocNode, data: dict) -> None:
    """Gene Set / Gene BMD section — per-(organ, sex) subsections."""
    role = genomics_role(node)
    _add_heading(doc, node.level, node.title, data)
    _body_style = _pstyle_or_default(doc, data, "body_para")
    intro = _add_paragraphs(doc, genomics_intro_paragraphs(node, data), style=_body_style)

    entries = genomics_entries(node, data)
    if not entries:
        if not intro:
            _add_pending(doc, f"Genomics data pending: {node.title}")
        return

    # One entry per organ (both sexes stacked in a single table — reference
    # Tables 9–12).  No per-sex subsection heading; the table's own **Male** /
    # **Female** separator rows delineate the sexes.
    for entry in entries:
        for item in genomics_content_plan(entry, role):
            _render_genomics_item(doc, entry, role, item)


def _render_genomics_item(doc: Document, entry: dict, role: str, item: dict) -> None:
    """One content item of a genomics (organ, sex) block."""
    part = item.get("part")
    if part == "narrative":
        _add_paragraphs(doc, entry.get("narrative") or [])
    elif part == "table":
        if role == "gene_set":
            _render_gene_set_table(doc, entry)
        else:
            _render_gene_table(doc, entry)
    elif part == "chart":
        chart = next(
            (c for c in (entry.get("charts") or []) if c.get("key") == item.get("chart_key")),
            None,
        )
        if chart:
            _add_chart_image(doc, chart)
    elif part == "descriptions":
        descriptions = (
            entry.get("go_descriptions") if role == "gene_set"
            else entry.get("gene_descriptions")
        ) or []
        _add_description_list(doc, descriptions)


def _render_figure(doc: Document, node: DocNode, data: dict) -> None:
    """A first-class figure node (ADR-0012): embed its lossless PNG + caption.

    The artifact is a payload dict at ``data[data_key]`` shaped like a chart
    payload (``{png_b64 | filename, caption}``) — so chart figures reuse the
    genomics chart pipeline's output verbatim and logo figures supply an image
    the same way.  Caption is ``node.caption`` (authored) or the payload's, prefixed
    with the positional ``Figure N.`` from node.figure_number.  A missing payload
    renders a visible pending note, never a silent gap."""
    payload = (data.get(node.data_key) if node.data_key else None) or {}
    png = payload.get("png_b64", "")
    if png.startswith("data:"):
        png = png.split(",", 1)[1]
    if png:
        try:
            doc.add_picture(BytesIO(base64.b64decode(png)), width=Inches(5.5))
        except Exception:
            _add_pending(doc, f"Figure image failed to decode: {node.title}")
    else:
        _add_pending(doc, f"Figure pending: {node.title}")
    text = node.caption or payload.get("caption") or node.title
    if text:
        label = f"Figure {node.figure_number}. " if node.figure_number else ""
        cap = doc.add_paragraph()
        run = cap.add_run(_clean(f"{label}{text}"))
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _render_gene_set_table(doc: Document, entry: dict) -> None:
    """Top-gene-sets table for one organ (both sexes stacked)."""
    rows = gene_set_table_rows(entry)
    if rows is None:
        _add_pending(doc, f"Top gene sets pending: {entry.get('organ','')}")
        return
    _booktabs_table(doc, list(GENE_SET_TABLE_HEADERS), rows,
                    caption=genomics_table_caption(entry), merged_label_rows=True)


def _render_gene_table(doc: Document, entry: dict) -> None:
    """Top-genes table for one organ (both sexes stacked)."""
    rows = gene_table_rows(entry)
    if rows is None:
        _add_pending(doc, f"Top genes pending: {entry.get('organ','')}")
        return
    _booktabs_table(doc, list(GENE_TABLE_HEADERS), rows,
                    caption=genomics_table_caption(entry), merged_label_rows=True)


def _add_description_list(doc: Document, descriptions) -> None:
    """A definition list of go-term / gene descriptions (bold term + text)."""
    for label, text in genomics_description_items(descriptions):
        para = doc.add_paragraph()
        if label:
            para.add_run(f"{_clean(label)}: ").bold = True
        para.add_run(_clean(text))


def _add_chart_image(doc: Document, chart: dict) -> None:
    """Embed a genomics chart PNG inline (decoded from its base64 payload)."""
    png = chart.get("png_b64", "")
    if not png:
        return
    if png.startswith("data:"):
        png = png.split(",", 1)[1]
    try:
        raw = base64.b64decode(png)
        doc.add_picture(BytesIO(raw), width=Inches(5.5))
    except Exception:
        return
    caption = genomics_chart_caption(chart)
    if caption:
        cap = doc.add_paragraph()
        run = cap.add_run(_clean(caption))
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _render_cover(doc: Document, node: DocNode, data: dict) -> None:
    """
    Branded page-1 cover — EXCLUDED on the Word surface (emits nothing).

    The reference cover (NIEHS Report 10, page 1) is hexagon artwork under a
    luminosity-masked gradient with Myriad Pro / Helvetica Neue text — Word can't
    embed that PDF page and the branded pixels aren't the fidelity target, so the
    docx skips the cover entirely and opens on the title page (the LaTeX surface
    keeps its tikz cover; see latex_generator._render_cover).  The handler stays
    registered so the dispatch table still covers every node type; it just
    contributes no content.
    """
    return


# The hardcoded per-role title-page spec — the measured reference default that
# stands in when a role has no configured style (so an empty config renders
# byte-identically to the pre-role behavior).  `title` roles = Arial Bold 20pt
# black; everything else = Times New Roman 12pt black; all centered.
#
# ``before``/``after`` are the per-role paragraph spacing IN POINTS, measured from
# the NTP example title page (examples/NIEHS-07…docx, the 1-NN styles).  Without
# them each line inherits Normal's 6pt-after, so the publisher block (five
# separate paragraphs) opened a 6pt gap between every line — the reported spacing
# bug.  The reference publisher lines are TIGHT (0 after); the deliberate gaps are
# `before` on publication_date (18), publisher_name (100 — the big drop that seats
# the block low on the page), and publisher_location (16), plus 6pt `after` on the
# title, report_number, and ISSN.  An explicit 0 must be set (not omitted) so it
# OVERRIDES Normal's inherited 6pt.
_TITLE_PAGE_ROLE_DEFAULTS: dict = {
    "report_title":           {"font": _HEADING_FONT, "size": 20, "bold": True,  "before": 0,   "after": 6},
    "report_number":          {"font": _BODY_FONT,    "size": 12, "bold": False, "before": 0,   "after": 6},
    "publication_date":       {"font": _BODY_FONT,    "size": 12, "bold": False, "before": 18,  "after": 0},
    "publisher_name":         {"font": _BODY_FONT,    "size": 12, "bold": False, "before": 100, "after": 0},
    "publication_institute":  {"font": _BODY_FONT,    "size": 12, "bold": False, "before": 0,   "after": 0},
    "publication_department": {"font": _BODY_FONT,    "size": 12, "bold": False, "before": 0,   "after": 0},
    "issn":                   {"font": _BODY_FONT,    "size": 12, "bold": False, "before": 0,   "after": 6},
    "publisher_location":     {"font": _BODY_FONT,    "size": 12, "bold": False, "before": 16,  "after": 0},
    # publisher_affiliation is the flat-fallback role (a layout with no per-line
    # roles); keep it tight so a fallback publisher block doesn't balloon either.
    "publisher_affiliation":  {"font": _BODY_FONT,    "size": 12, "bold": False, "before": 0,   "after": 0},
}
_TITLE_PAGE_META_DEFAULT = {"font": _BODY_FONT, "size": 12, "bold": False, "before": 0, "after": 6}


def _resolve_title_page_role(layout_cfg: "dict | None", role: str) -> dict:
    """The configured style for a title-page role, resolved in precedence order:
    ``defaults`` ← ``types["title-page"]`` (a page-wide baseline) ←
    ``title_page[role]`` (the per-role override).  Empty when nothing configured
    (the handler then keeps the measured hardcoded spec).

    NB the handler owns ALL title-page styling: the generic per-node overlay in
    `_walk_docx_tree` is SKIPPED for this node (see `_visit`), so this resolver
    must fold in the node-level ``types``/``instances`` layers itself — otherwise
    a page-wide `types["title-page"]` style would be lost."""
    if not layout_cfg:
        return {}
    style = dict(layout_cfg.get("defaults") or {})
    style.update((layout_cfg.get("types") or {}).get("title-page") or {})
    style.update((layout_cfg.get("instances") or {}).get("title-page") or {})
    style.update((layout_cfg.get("title_page") or {}).get(role) or {})
    return style


def _render_title_page(doc: Document, node: DocNode, data: dict) -> None:
    """
    Inner title page (reference page 2) — a centered typographic block.  Each
    semantic ROLE (report_title, report_number, publication_date, publisher_name,
    publisher_affiliation, issn, …) is emitted as ONE paragraph (multi-line roles
    like the title use internal line breaks, matching the reference's single
    title paragraph — not one paragraph per line).

    Styling precedence per role: the measured hardcoded reference spec
    (`_TITLE_PAGE_ROLE_DEFAULTS`) is applied first, then the configured
    ``styles.title_page[role]`` (resolved over ``defaults``) is overlaid via the
    shared `_apply_paragraph_style`, so any key the config omits keeps its
    reference value (ADR-0006 no-drift: an empty config renders unchanged).

    Text comes from the SAME cover_layouts builder the wording is single-sourced
    through, so the surfaces can't drift on wording.  Falls back to the flat
    title/publisher builders when a layout supplies no role builder.
    """
    layout = get_cover_layout(node.subtype)
    layout_cfg = data.get("layout_style")

    blocks = layout.title_page_blocks(data) if layout.title_page_blocks else None
    if blocks is None:
        # Legacy fallback: a layout without the role builder → old flat behavior.
        blocks = [("report_title", layout.title_builder(data))]
        for meta in (data.get("report_number", ""), data.get("report_date", "")):
            if meta:
                blocks.append(("report_number", [meta]))
        blocks.append(("publisher_affiliation", layout.publisher_builder(data)))

    for role, lines in blocks:
        lines = [_clean(ln) for ln in lines if ln]
        if not lines:
            continue
        para = doc.add_paragraph()
        run = para.add_run(lines[0])
        for extra in lines[1:]:
            run.add_break(WD_BREAK.LINE)
            run = para.add_run(extra)
        # Role-driven path (ADR-0010 Phase 2): when a vocabulary is active, tag the
        # paragraph with the role's NATIVE Word style — the title inherits the NTP
        # 1-NN typography (incl. the neutral single line spacing) with NO hardcoded
        # defaults.  Otherwise fall back to the measured reference defaults +
        # optional config overlay (the pre-vocabulary look).
        style_name = _role_style_name(data, role)
        if style_name:
            try:
                para.style = doc.styles[style_name]
            except KeyError:
                style_name = None
        if not style_name:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            base = _TITLE_PAGE_ROLE_DEFAULTS.get(role, _TITLE_PAGE_META_DEFAULT)
            for r in para.runs:
                r.font.name = base["font"]
                r.font.size = Pt(base["size"])
                r.font.bold = base["bold"]
                r.font.color.rgb = _BLACK
            # Per-role paragraph spacing (points), measured from the NTP reference.
            # Set explicitly — including 0 — so it overrides Normal's inherited
            # 6pt-after; otherwise the multi-paragraph publisher block gaps out.
            pf = para.paragraph_format
            if "before" in base:
                pf.space_before = Pt(base["before"])
            if "after" in base:
                pf.space_after = Pt(base["after"])
        # A styles.title_page per-role config override ALWAYS applies on top —
        # on the base/vocab path it overlays the NTP named style, on the legacy
        # path it overlays the measured defaults above.  So a user can still
        # re-font/re-size an individual title-page role regardless of the base.
        role_style = _resolve_title_page_role(layout_cfg, role)
        if role_style:
            _apply_paragraph_style(para, role_style)

    # The title page is a self-contained front page: always end it so the next
    # node (Foreword) starts fresh, regardless of the active tree or whether a
    # page-break node was authored between them.  The reference separates the
    # title block from the Foreword with a section break; a page break is the
    # equivalent here (the front-matter section break at the body boundary is
    # emitted separately by generate_docx).  Previously this separation happened
    # only by accident, via the body-heading style's break_before — which
    # vanished once the Foreword correctly used the (no-break) 1-21 title style.
    _add_page_break(doc)


def _freeform_text(node: DocNode) -> str:
    """
    Plain text for a freeform node on the Word surface.  There is no docx entry
    in resolved_content ({latex, html}); fall back to the html body with tags
    stripped (readable prose), else a pending note naming the authored surface.
    """
    resolved = node.resolved_content or {}
    html = resolved.get("html")
    if html:
        text = _TAG_RE.sub("", html)
        return re.sub(r"\s+\n", "\n", text).strip()
    rep = node.representation or "latex"
    return _freeform_pending_note(rep, "html")


def _render_freeform_page(doc: Document, node: DocNode, data: dict) -> None:
    """Freeform authored page — forces its own page, then heading + body text."""
    _add_page_break(doc)
    if node.title:
        _add_heading(doc, node.level, node.title, data)
    _add_paragraphs(doc, _freeform_text(node).split("\n\n"))


def _render_freeform_block(doc: Document, node: DocNode, data: dict) -> None:
    """Freeform authored block — inline, no forced page break."""
    if node.title:
        _add_heading(doc, node.level, node.title, data)
    _add_paragraphs(doc, _freeform_text(node).split("\n\n"))


def _render_page_break(doc: Document, node: DocNode, data: dict) -> None:
    """An explicit author-placed page break."""
    _add_page_break(doc)


def _render_unimplemented(doc: Document, node: DocNode, data: dict) -> None:
    """Catch-all — heading (if any) + a visible pending placeholder."""
    _add_heading(doc, node.level, node.title, data)
    _add_pending(doc, f"Section pending: {node.node_type} rendering not yet implemented")


def _format_dose_label(dose, unit: str) -> str:
    """Format a dose value with its unit (the Word/plain-space form)."""
    if dose == 0 or dose == 0.0:
        return f"0 {unit}"
    if isinstance(dose, float) and dose.is_integer():
        return f"{int(dose)} {unit}"
    return f"{dose} {unit}"


# ---------------------------------------------------------------------------
# Dispatch registry — validated against the canonical set at import (ADR-0006 #3)
# ---------------------------------------------------------------------------

_DISPATCH: dict[str, object] = {
    "cover":               _render_cover,
    "title-page":          _render_title_page,
    "front-matter":        _render_front_matter,
    "narrative":           _render_narrative,
    "heading-only":        _render_heading_only,
    "appendix":            _render_appendix,
    "tables-list":         _render_tables_list,
    "toc":                 _render_toc,
    "narrative+tables":    _render_narrative_tables,
    "table":               _render_apical_table,
    "incidence-table":     _render_incidence_table,
    "figure":              _render_figure,
    "sample-counts-table": _render_sample_counts_table,
    "bmd-summary":         _render_bmd_summary,
    "genomics-section":    _render_genomics_section,
    "freeform-page":       _render_freeform_page,
    "freeform-block":      _render_freeform_block,
    "page-break":          _render_page_break,
}

# Fail loudly at import if this table drifts from the canonical registry — the
# same guarantee the HTML and LaTeX surfaces get (ADR-0006 #3).  Word implements
# every renderable type, so no omissions are allowed.
assert_dispatch_covers(_DISPATCH, renderer="Word/OOXML")


# ---------------------------------------------------------------------------
# Per-node layout styling (the docx twin of _layout_to_css_props / _layout_to_latex)
# ---------------------------------------------------------------------------

def _length_to_pt(value) -> "float | None":
    """A resolved-style length ('6pt', '0.5in') → points; None if unparseable or
    in a relative unit (em/ex) we can't resolve without the current font size."""
    m = _LENGTH_RE.match(str(value or ""))
    if not m:
        return None
    factor = _PT_PER_UNIT.get(m.group(2))
    return float(m.group(1)) * factor if factor is not None else None


def _resolve_font_name(style: dict) -> "str | None":
    """The concrete Word font for a resolved style: the literal `font` wins, else
    map `font_family` (serif/sans/mono) through _DOCX_FONT_BY_FAMILY.  See the
    precedence note in layout_style.LAYOUT_KEY_SCHEMA."""
    name = (style.get("font") or "").strip()
    if name:
        return name
    return _DOCX_FONT_BY_FAMILY.get(style.get("font_family"))


def _apply_run_style(run, style: dict) -> None:
    """Apply the character-level part of a resolved style to one run (delegates to
    _apply_font_style on the run's font)."""
    _apply_font_style(run.font, style)


def _apply_font_style(font, style: dict) -> None:
    """Apply the character-level part of a resolved style to a Font object — works
    for both a run's font and a named style's font (both expose the same Font API
    + a `_element` carrying an rPr).  So the same character logic bakes into a
    style DEFINITION (Phase 2) or overlays a run (legacy path)."""
    name = _resolve_font_name(style)
    if name:
        font.name = name
    size = _length_to_pt(style.get("font_size"))
    if size:
        font.size = Pt(size)
    weight = style.get("weight")
    if weight == "bold":
        font.bold = True
    elif weight == "normal":
        font.bold = False
    st = style.get("style")
    if st == "italic":
        font.italic = True
    elif st == "normal":
        font.italic = False
    tt = style.get("text_transform")
    if tt == "uppercase":
        font.all_caps = True   # w:caps — a DISPLAY transform (text unchanged)
    elif tt == "none":
        font.all_caps = False
    color = style.get("color")
    if isinstance(color, str) and color.startswith("#"):
        h = color.lstrip("#")
        if len(h) == 3:
            h = "".join(c * 2 for c in h)
        if len(h) == 6:
            font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    # letter_spacing → rPr <w:spacing w:val="TWIPS"> (character spacing, in twips =
    # pt*20).  python-docx's Font exposes no `spacing`, so set it on the rPr
    # directly.  Font._element is the run (rPr host) or the style element (both
    # answer get_or_add_rPr).  Absolute units only (em/ex → _length_to_pt None).
    ls_pt = _length_to_pt(style.get("letter_spacing"))
    if ls_pt is not None:
        rpr = font._element.get_or_add_rPr()
        spacing = rpr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            rpr.append(spacing)
        spacing.set(qn("w:val"), str(int(round(ls_pt * 20))))


def _apply_outline_level(style_element, style: dict) -> None:
    """Stamp <w:outlineLvl w:val="N"> onto a style's pPr when the resolved style
    carries ``outline_level`` (0-based).  Outline level is what a Word TOC field
    (and the Navigation pane) collects by — the NTP role heading styles carry no
    outlineLvl of their own, so without this the field on the vocabulary path
    collects nothing.  A no-op when the key is absent (body/caption styles).
    python-docx exposes no outline-level API on a style, so edit the pPr directly.
    """
    lvl = style.get("outline_level")
    if not isinstance(lvl, (int, float)) or isinstance(lvl, bool):
        return
    pPr = style_element.get_or_add_pPr()
    ol = pPr.find(qn("w:outlineLvl"))
    if ol is None:
        ol = OxmlElement("w:outlineLvl")
        pPr.append(ol)
    ol.set(qn("w:val"), str(int(lvl)))


def _apply_style_props_to_style(word_style, style: dict) -> None:
    """Apply a resolved style dict to a Word STYLE object (not a paragraph): its
    paragraph_format + its font.  Used to bake a vocabulary type's OWN delta into
    a named <w:style> (ADR-0010 Phase 2), so the style DEFINITION carries the
    property and every paragraph tagged with it inherits — the native Word model,
    unlike the per-paragraph overlay _apply_paragraph_style does for the legacy
    path.  Only the character part differs (a style has one .font, not runs)."""
    if not style:
        return
    pf = word_style.paragraph_format
    align = style.get("align")
    if align:
        pf.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "center": WD_ALIGN_PARAGRAPH.CENTER, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(align)
    lh = style.get("line_height")
    if isinstance(lh, (int, float)) and not isinstance(lh, bool):
        pf.line_spacing = lh
    sb = _length_to_pt(style.get("space_before"))
    if sb is not None:
        pf.space_before = Pt(sb)
    sa = _length_to_pt(style.get("space_after"))
    if sa is not None:
        pf.space_after = Pt(sa)
    indent = _length_to_pt(style.get("first_line_indent"))
    if indent is not None:
        pf.first_line_indent = Pt(indent)
    if style.get("keep_together") is True:
        pf.keep_together = True
    if style.get("break_before") == "page":
        pf.page_break_before = True
    _apply_outline_level(word_style.element, style)
    _apply_font_style(word_style.font, style)


def _apply_paragraph_style(paragraph, style: dict) -> None:
    """Apply the paragraph-level part of a resolved style to one paragraph, and
    the character-level part to each of its runs (so inheritance is explicit)."""
    pf = paragraph.paragraph_format
    align = style.get("align")
    if align:
        pf.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "center": WD_ALIGN_PARAGRAPH.CENTER, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(align)
    lh = style.get("line_height")
    if isinstance(lh, (int, float)) and not isinstance(lh, bool):
        pf.line_spacing = lh
    sb = _length_to_pt(style.get("space_before"))
    if sb is not None:
        pf.space_before = Pt(sb)
    sa = _length_to_pt(style.get("space_after"))
    if sa is not None:
        pf.space_after = Pt(sa)
    indent = _length_to_pt(style.get("first_line_indent"))
    if indent is not None:
        pf.first_line_indent = Pt(indent)
    if style.get("keep_together") is True:
        pf.keep_together = True
    # NB: break_before / break_after are NODE-level flow, applied once at the node
    # boundary by _layout_to_docx (not here) — see the note there.  Applying them
    # per paragraph would break before/after EVERY paragraph of a multi-paragraph
    # node, diverging from HTML (one wrapping div) and LaTeX (one \clearpage).
    for run in paragraph.runs:
        _apply_run_style(run, style)


def _layout_to_docx(paragraphs, style: dict) -> None:
    r"""
    Apply a resolved abstract style spec to a run of paragraphs — the docx twin
    of html_generator._layout_to_css_props and latex_generator._layout_to_latex.

    Unlike those string-wrapping surfaces, Word is an object model: the node's
    handler has ALREADY appended its paragraphs, so this overlays the per-node
    ``types``/``instances`` styling onto them (the base ``Normal``/``Heading``
    style from _build_style_skeleton is the baseline; this wins over it, mirroring
    the other two surfaces' _wrap_style layering).  Empty style ⇒ no-op, so a
    style-less document is byte-identical to the pre-feature output.

    ``break_before`` / ``break_after`` are NODE-level flow and are applied ONCE
    here at the node boundary — before the first paragraph, after the last —
    matching HTML's single wrapping div and LaTeX's single \clearpage.  Word has
    no paragraph "page-break-after" property (OOXML only offers pageBreakBefore),
    so an after-break is emitted as a trailing page-break run; this is why the
    extractor can read ``break_before`` from a style but never ``break_after``.
    """
    if not style:
        return
    paragraphs = list(paragraphs)
    for paragraph in paragraphs:
        _apply_paragraph_style(paragraph, style)
    if not paragraphs:
        return
    if style.get("break_before") == "page":
        paragraphs[0].paragraph_format.page_break_before = True
    if style.get("break_after") == "page":
        paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Tree walk + document skeleton
# ---------------------------------------------------------------------------

def _effective_landscape(node: DocNode, data: dict) -> bool:
    """The node's EFFECTIVE orientation, as the shared FLOW-axis decision.

    Delegates to render_capabilities.landscape_requested so all three surfaces
    resolve orientation in exactly ONE place (override > template default >
    portrait, gated on the type's orientable capability).  This is the abstract
    per-node FLOW value; the Word adapter's job is to LOWER it into sections."""
    return landscape_requested(
        node.node_type, node.id, data.get("orientations"), default=node.orientation
    )


def _open_flow_section(doc: Document, document: dict | None, landscape: bool):
    """Open a new page-geometry section for a FLOW (orientation) transition.

    Word entangles orientation with the section (`w:sectPr`) construct: there is
    no per-paragraph orientation, so flipping orientation REQUIRES a new section.
    This is the lowering from our orthogonal, per-node FLOW axis to Word's
    per-region one — a landscape "island" is a section bracketed by portrait
    sections (exactly how the NIEHS-10 reference encodes its two wide-table
    spreads).

    The section is REGION-TRANSPARENT: it is created but NOT run through
    _configure_section, so it does not declare its own footer/header or
    <w:pgNumType>.  A fresh section's header/footer start LINKED to the previous
    one and it carries no page-number restart, so the body's running footer and
    arabic counter FLOW THROUGH the island unbroken — the reference's landscape
    sections likewise omit footerReference and pgNumType.  (Orientation and page
    numbering are independent axes that Word happens to store in the same
    element; keeping the FLOW section transparent is what keeps them orthogonal.)

    Geometry: python-docx does NOT auto-swap width/height on orientation change,
    and add_section CLONES the previous section's pgSz (including its `orient`
    attribute), so BOTH the orientation flag and the dimensions must be set
    explicitly on every transition — including the portrait RESUME, whose clone
    would otherwise inherit the landscape orientation with portrait dimensions.
    """
    document = document or {}
    pw = _emu_or_default(document, "page_width", _PAGE_WIDTH)
    ph = _emu_or_default(document, "page_height", _PAGE_HEIGHT)
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = ph, pw
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = pw, ph
    # No cover here, so no distinct first-page header (matches the body sections).
    section.different_first_page_header_footer = False
    # add_section CLONES the previous section's whole sectPr, so it drags along a
    # <w:pgNumType> (the styles base carries one, w:start="1").  Left in place,
    # this FLOW section would RESTART page numbering to 1 mid-body — the opposite
    # of region-transparency.  Strip it so the section inherits the body counter;
    # the reference's landscape sections carry no pgNumType at all.
    pgNumType = section._sectPr.find(qn("w:pgNumType"))
    if pgNumType is not None:
        section._sectPr.remove(pgNumType)
    return section


def _walk_docx_tree(
    doc: Document, node: DocNode, data: dict, flow: dict, document: dict | None = None
) -> None:
    """
    Render a node and its descendants into the Document in document order.

    The Word analogue of html_generator._walk_html / latex_generator._walk_latex,
    but it can't reuse render_common.walk_emit: that skeleton accumulates markup
    STRINGS, whereas Word handlers mutate the Document.  So this runs the shared
    walk_tree primitive directly with a _visit callback that dispatches to the
    (doc, node, data) handlers — same pre-order traversal, object-model emit.

    Per-node layout styling (ADR-0006 parity with the other surfaces): the
    resolved ``types``/``instances`` spec is applied to exactly the paragraphs a
    node's handler added, by snapshotting the paragraph count before/after the
    handler runs.  The DECISION (resolved spec) is shared via
    data["layout_style"]; only this object-model application is surface-specific.

    FLOW lowering: ``flow`` is a mutable {"landscape", "in_body"} state threaded
    across the whole body walk (it persists between top-level nodes).  When a
    node's EFFECTIVE orientation differs from the current section's, a new
    section is opened BEFORE the node renders — coalescing a maximal contiguous
    run of same-orientation nodes into ONE section (per-node FLOW → per-region
    Word section).  The transition runs before the paragraph-count snapshot so
    the empty paragraph add_section injects (which carries the CLOSING section's
    sectPr) is not misattributed to this node's layout styling.  Only active in
    the body region; front matter carries no orientable nodes.
    """
    layout_cfg = data.get("layout_style")

    def _visit(n: DocNode) -> None:
        # FLOW axis → Word sections: open/resume a section when orientation flips.
        if flow.get("in_body"):
            want = _effective_landscape(n, data)
            if want != flow["landscape"]:
                _open_flow_section(doc, document, landscape=want)
                flow["landscape"] = want
        before = len(doc.paragraphs)
        handler = _DISPATCH.get(n.node_type, _render_unimplemented)
        handler(doc, n, data)
        # The title-page handler applies its own PER-ROLE styling (it folds in the
        # node-level types/instances layers itself); skip the generic uniform
        # overlay so it doesn't clobber the per-role fonts/sizes with one style.
        if layout_cfg and n.node_type != "title-page":
            style = resolve_layout_style(layout_cfg, n.node_type, n.id)
            if style:
                _layout_to_docx(doc.paragraphs[before:], style)

    walk_tree([node], _visit)


def _set_section_page_numbering(section, fmt: str, start: int) -> None:
    """
    Set a section's <w:pgNumType> format (lowerRoman / decimal) + start.

    <w:pgNumType> must sit at its schema-mandated position in CT_SectPr (after
    pgSz/pgMar, before cols) — a bare append lands it after docGrid, which is
    out of sequence and gets silently dropped on save.  We insert it before the
    first of the elements the schema says follow it, so the ordering is valid
    regardless of which of those a given section happens to carry.
    """
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        _insert_in_schema_order(
            sectPr, pgNumType,
            successors=("w:cols", "w:formProt", "w:vAlign", "w:noEndnote",
                        "w:titlePg", "w:textDirection", "w:bidi", "w:rtlGutter",
                        "w:docGrid", "w:printerSettings", "w:sectPrChange"),
        )
    pgNumType.set(qn("w:fmt"), fmt)
    pgNumType.set(qn("w:start"), str(start))


def _insert_in_schema_order(parent, element, *, successors: tuple[str, ...]) -> None:
    """
    Insert ``element`` before the first existing child whose tag is in
    ``successors`` (the elements the schema requires to come AFTER it), else
    append.  Keeps a hand-built element valid against the OOXML sequence.
    """
    for tag in successors:
        found = parent.find(qn(tag))
        if found is not None:
            found.addprevious(element)
            return
    parent.append(element)


def _mark_fields_dirty(doc: Document) -> None:
    """Set <w:updateFields w:val="true"/> in settings.xml so Word/LibreOffice
    recompute every field on open — the TOC field then populates (page numbers +
    dot leaders + links) without a manual "update field".  LibreOffice applies
    this silently; Word shows a one-time "update fields?" prompt.  Idempotent.

    <w:updateFields> sits LATE in the CT_Settings sequence (after the XML/preview
    settings, before compat/rsids/mathPr/...).  A bare prepend or append lands it
    out of order, and an out-of-sequence settings child is SILENTLY DROPPED on
    load — the same trap _set_section_page_numbering documents for pgNumType — so
    the field would never update.  Insert it before the first of its schema
    successors instead."""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is not None:
        return
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    _insert_in_schema_order(
        settings, el,
        successors=("w:hdrShapeDefaults", "w:footnotePr", "w:endnotePr",
                    "w:compat", "w:rsids", "w:mathPr", "w:uiCompat97To2003",
                    "w:attachedSchema", "w:themeFontLang", "w:clrSchemeMapping",
                    "w:doNotIncludeSubdocsInStats", "w:doNotAutoCompressPictures",
                    "w:forceUpgrade", "w:captions", "w:readModeInkLockDown",
                    "w:smartTagType", "w:shapeDefaults", "w:doNotEmbedSmartTags",
                    "w:decimalSymbol", "w:listSeparator", "w:docId"),
    )


def _add_page_number_field(paragraph) -> None:
    """Insert a PAGE field into a (footer) paragraph, centered by the caller."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _emu_or_default(document: dict, key: str, default):
    """A document-level length in EMU (via docx Pt/Inches types) or the default.
    Lengths are stored as strings ('12pt', '1in'); parse to pt then to a
    Length."""
    pt = _length_to_pt(document.get(key)) if document else None
    return Pt(pt) if pt is not None else default


def _load_active_vocabulary(data: dict):
    """Load the vocabulary named by ``data['vocabulary']`` (a vocab/<name>.yaml
    stem), or None when unset — the opt-in switch for role-driven styling.  A load
    error is swallowed to None (fall back to legacy styling) rather than breaking
    generation; the vocabulary is an enhancement, not a hard dependency."""
    name = (data.get("vocabulary") or "").strip()
    if not name:
        return None
    try:
        import vocabulary as _vocab
        return _vocab.load_vocabulary(name)
    except Exception:
        return None


def _role_style_name(data: dict, role: str) -> "str | None":
    """The Word style name a vocabulary role maps to (its docx binding), or None
    when no vocabulary is active or the role is unknown — the switch that makes a
    handler apply a native pStyle only on the role-driven path, falling back to
    its legacy styling otherwise."""
    vocab = data.get("_vocabulary")
    if vocab is None or vocab.get(role) is None:
        return None
    import vocabulary as _vocab
    return _vocab.resolve_bindings(vocab, role)["docx"]


def _pstyle_or_default(doc: Document, data: dict, role: str) -> "str | None":
    """Resolve a vocabulary ``role`` to a Word style NAME that actually EXISTS in
    ``doc``, or None when it can't (no vocab, unknown role, or the bound style is
    absent from the style base).  Callers apply the returned name as a pStyle and
    treat None as "use the default (Normal)" — so a missing style degrades to
    plain body text instead of raising KeyError (the old hardcoded-name trap)."""
    name = _role_style_name(data, role)
    if name and name in {s.name for s in doc.styles}:
        return name
    return None


def _build_vocabulary_styles(doc: Document, vocab) -> None:
    """Emit the vocabulary's type graph as native Word paragraph styles (ADR-0010
    Phase 2): one <w:style> per type, styleId/name = its docx binding, basedOn =
    the specialization parent's binding, properties = the type's OWN style delta.
    The resulting styles.xml mirrors the NTP basedOn graph, so applying a role by
    pStyle resolves through the same inheritance the reference uses — and a
    co-author opening the docx sees the real named-style palette.

    Built in specialization order (parents before children) so base_style can be
    assigned.  A type whose docx binding is an EXISTING style (Normal, Heading 1)
    updates that style in place rather than adding a duplicate."""
    import vocabulary as _vocab
    from docx.enum.style import WD_STYLE_TYPE

    styles = doc.styles

    # Topological order: a type appears after its `specializes` parent.  The
    # graph is a shallow forest (roots → NTP roles), so a simple resolved-set
    # sweep terminates quickly.
    remaining = dict(vocab.types)
    ordered: list[str] = []
    placed: set[str] = set()
    while remaining:
        progressed = False
        for name, rec in list(remaining.items()):
            parent = rec.specializes
            if parent is None or parent in placed or parent not in vocab.types:
                ordered.append(name)
                placed.add(name)
                del remaining[name]
                progressed = True
        if not progressed:  # defensive: a cycle (validation should prevent this)
            ordered.extend(remaining)
            break

    for name in ordered:
        rec = vocab.types[name]
        bindings = _vocab.resolve_bindings(vocab, name)
        style_name = bindings["docx"]
        # Reuse an existing Word style of that name, else create a paragraph style.
        try:
            style = styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        # basedOn = the parent type's docx style (skip if the parent isn't a real
        # style in the doc, e.g. a pure grouping root whose binding we didn't emit).
        # CRITICAL: skip when the parent binds to the SAME Word style name — a
        # curated alias (e.g. publisher_location) may reuse its concrete parent's
        # Word style; both vocab types then map to one physical <w:style>, and
        # setting basedOn to itself is a self-reference Word can't resolve (the
        # chain breaks → left/default rendering).  Same-name ⇒ same style ⇒ no edge.
        if rec.specializes and rec.specializes in vocab.types:
            parent_name = _vocab.resolve_bindings(vocab, rec.specializes)["docx"]
            if parent_name != style_name:
                try:
                    style.base_style = styles[parent_name]
                except KeyError:
                    pass
        # Apply ONLY this type's own delta (inheritance carries the rest), so the
        # emitted style mirrors the NTP delta-based definitions.
        if rec.style:
            _apply_style_props_to_style(style, rec.style)


def _neutralize_docdefaults_spacing(styles) -> None:
    """Remove python-docx's default docDefaults paragraph spacing entirely, so the
    document root is spacing-neutral like the NTP reference.

    A fresh Document() ships docDefaults pPr <w:spacing w:line="276" (1.15x)
    w:after="200" (10pt)> — the modern-Office default.  The NTP reference's
    docDefaults has NO spacing element at all (verified): every gap comes from the
    individual 1-NN/0-NN STYLES, never a document-wide default.  Leaving the
    default in place means (a) the title inherits 1.15x line spacing, and (b) a
    paragraph without its own space_after inherits a 10pt gap — which is why a
    hand-inserted paragraph on the title page opened a huge gap.  We drop the
    whole <w:spacing> element to match the reference; each style's own
    space_before/after then governs, and single line spacing applies by default.
    python-docx exposes no docDefaults API, so we edit the element directly.
    Idempotent + safe if the element is absent."""
    el = styles.element
    ppr = el.find(
        "/".join(qn(t) for t in ("w:docDefaults", "w:pPrDefault", "w:pPr"))
    )
    if ppr is None:
        return
    spacing = ppr.find(qn("w:spacing"))
    if spacing is not None:
        ppr.remove(spacing)


def _build_style_skeleton(doc: Document, document: dict | None = None) -> None:
    """
    Bind the base typography onto the document's Word styles BEFORE any content
    is walked, replacing Word's default theme (Calibri/Aptos).

    All handlers create paragraphs via the ``Normal`` / ``Heading 1-3`` styles,
    so restyling those here cascades to the whole document — the base look lives
    in ONE place rather than being set per run.  Fonts are named by FAMILY (the
    Font.name setter writes both w:ascii and w:hAnsi runProps), so Word uses the
    installed font regardless of its theme's minor/major font.

    ``document`` is the optional resolved document-level style block
    (styles.document); its ``default_font`` / ``default_font_size`` /
    ``header_font`` override the measured constants when present, so the base
    body/header fonts become DATA.  Absent ⇒ the measured reference constants.
    """
    document = document or {}
    body_font = (document.get("default_font") or "").strip() or _BODY_FONT
    body_size = _length_to_pt(document.get("default_font_size")) or _BODY_PT
    header_font = (document.get("header_font") or "").strip() or _BODY_FONT
    header_size = _length_to_pt(document.get("header_font_size")) or _HEADER_PT

    styles = doc.styles

    # Neutralize python-docx's non-reference docDefaults (ADR-0010): a fresh
    # Document() ships docDefaults pPr <w:spacing w:line="276" lineRule="auto">
    # (1.15x) + after=200, the modern-Office default.  The NTP reference sets NO
    # line spacing at the document root, so every part — including the title,
    # which inherits the root, not Normal — renders SINGLE.  Our 1.15x root is
    # exactly why the generated title looked mis-spaced.  Strip the docDefaults
    # paragraph spacing so the neutral single default (vocab/base.yaml `text`)
    # applies; Normal below still sets its own space_after for body paragraphs.
    _neutralize_docdefaults_spacing(styles)

    normal = styles["Normal"]
    normal.font.name = body_font
    normal.font.size = Pt(body_size)
    normal.font.color.rgb = _BLACK
    npf = normal.paragraph_format
    # Times' natural single spacing gives the reference's ~13.8pt line pitch;
    # block paragraphs (no first-line indent) with a little space after.
    npf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    npf.space_before = Pt(0)
    npf.space_after = Pt(6)
    npf.first_line_indent = Pt(0)

    for level, size in _HEADING_PT.items():
        style = styles[f"Heading {level}"]
        style.font.name = _HEADING_FONT
        style.font.bold = True
        style.font.size = Pt(size)
        # Word's built-in headings are blue — the reference headings are black.
        style.font.color.rgb = _BLACK
        hpf = style.paragraph_format
        hpf.keep_with_next = True
        hpf.space_before = Pt(12)
        hpf.space_after = Pt(4)

    # The running header rides on the built-in "Header" style.
    header = styles["Header"]
    header.font.name = header_font
    header.font.size = Pt(header_size)


# Max characters per running-header line.  The header runs ~12pt across the full
# ~6.5" text block, so it fits far more per line than the 20pt title; we wrap only
# to avoid Word's arbitrary mid-phrase break.  The reference breaks its ~85-char
# header into two lines after "…Study of", so ~50 puts the break at a comparable
# semantic point without overflowing.
_HEADER_MAX_CHARS = 50


def _set_header_text(header_para, text: str) -> None:
    """Set the running-header text, width-wrapped into lines joined by soft breaks
    so it breaks at word boundaries rather than auto-wrapping mid-phrase (the
    reported header-wrap bug).  Reuses cover_layouts._wrap_words for the same
    greedy packing the title uses.  A short header stays one line."""
    from cover_layouts import _wrap_words

    for run in list(header_para.runs):        # clear any existing runs
        run._element.getparent().remove(run._element)
    lines = _wrap_words(text, _HEADER_MAX_CHARS) or [text]
    run = header_para.add_run(lines[0])
    for extra in lines[1:]:
        run.add_break(WD_BREAK.LINE)
        run = header_para.add_run(extra)


def _configure_section(section, running_header: str, document: dict | None = None) -> None:
    """US-Letter geometry + a running header + a centered page-number footer.

    ``document`` (the resolved styles.document block) overrides the page size,
    margins, and header distance when present; absent ⇒ the measured constants."""
    document = document or {}
    section.page_width = _emu_or_default(document, "page_width", _PAGE_WIDTH)
    section.page_height = _emu_or_default(document, "page_height", _PAGE_HEIGHT)
    section.top_margin = _emu_or_default(document, "margin_top", _MARGIN)
    section.bottom_margin = _emu_or_default(document, "margin_bottom", _MARGIN)
    section.left_margin = _emu_or_default(document, "margin_left", _MARGIN)
    section.right_margin = _emu_or_default(document, "margin_right", _MARGIN)
    hdr = _emu_or_default(document, "header_distance", _HEADER_FOOTER_DISTANCE)
    section.header_distance = hdr
    section.footer_distance = hdr

    header_para = section.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Break the running header at word boundaries so it does not auto-wrap at an
    # arbitrary point in the text-block width (the reference inserts an explicit
    # break after "…Study of" rather than letting Word wrap).  Width-wrap into
    # lines with soft breaks; one line renders unchanged.
    _set_header_text(header_para, _clean(running_header))

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # A later section's footer is LINKED to the previous one (they share the same
    # footer element), so configuring both sections would append a SECOND PAGE
    # field to the one shared footer — the field then renders twice ("iii" → the
    # reported "iiii").  The per-section roman/decimal split comes from pgNumType,
    # not the footer, so one shared PAGE field is correct; only add it if absent.
    if not footer_para._p.findall(".//" + qn("w:instrText")):
        _add_page_number_field(footer_para)


# The styles-only base docx (built by _build_docx_base.py from the NIEHS-10
# reference): a full 386-style NTP library with an empty body.  Opening it as the
# base — instead of a blank Document() whose python-docx defaults we'd re-derive
# in Python — makes styles/docDefaults/numbering/theme/sectPr come from the
# reference VERBATIM.  Absent ⇒ fall back to a blank Document() + the
# programmatic style skeleton (the legacy path), so generation never hard-breaks.
_DOCX_BASE_PATH = Path(__file__).with_name("assets") / "templates" / "niehs-10-base.docx"

# The vocabulary applied by default on the docx surface when the style base is
# used: role → NTP named style (3-02a_Head1_NoNumber, 0-03_Paragraph, …).  The
# base already CONTAINS these styles, so we only need the vocab for role→name
# resolution in the content path, not to build styles.
_DEFAULT_DOCX_VOCAB = "ntp-report"


def _open_base_document() -> "tuple[Document, bool]":
    """Open the NTP styles-only base docx; return (doc, used_base).

    used_base=False when the asset is absent — the caller then builds the legacy
    programmatic style skeleton on the blank doc instead.  The base's single
    placeholder body paragraph is cleared here so the content walk starts clean;
    its trailing sectPr (page geometry / docGrid) is preserved."""
    if _DOCX_BASE_PATH.exists():
        try:
            doc = Document(str(_DOCX_BASE_PATH))
            _clear_base_body(doc)
            return doc, True
        except Exception:
            # A corrupt/unreadable base must never break generation — fall back.
            pass
    return Document(), False


def _clear_base_body(doc: Document) -> None:
    """Remove the base's placeholder body content, preserving the trailing sectPr
    (page geometry lives there) so the section is well-formed for the walk."""
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is sectPr:
            continue
        body.remove(child)


def _apply_document_font_overrides(doc: Document, document: dict) -> None:
    """On the template-base path, honor the styles.document block's base-FONT
    overrides (default_font / default_font_size / header_font / header_font_size)
    by setting them on Normal + Header.  The base already supplies the reference
    fonts, so this ONLY runs when the caller explicitly overrides them — geometry
    overrides are handled separately by _configure_section.  A no-op when the
    document block names no fonts (the common case), so the base's own fonts win."""
    if not document:
        return
    body_font = (document.get("default_font") or "").strip()
    body_size = _length_to_pt(document.get("default_font_size"))
    header_font = (document.get("header_font") or "").strip()
    header_size = _length_to_pt(document.get("header_font_size"))
    if body_font or body_size:
        normal = doc.styles["Normal"]
        if body_font:
            normal.font.name = body_font
        if body_size:
            normal.font.size = Pt(body_size)
    if header_font or header_size:
        try:
            header = doc.styles["Header"]
        except KeyError:
            header = None
        if header is not None:
            if header_font:
                header.font.name = header_font
            if header_size:
                header.font.size = Pt(header_size)


def generate_docx(data: dict, tree: "list | None" = None) -> bytes:
    """
    Walk the document tree + data and produce a complete .docx as bytes.

    Same data dict the other two surfaces consume (report_data.marshal_export_data
    / latex_export.load_session_data).  Front matter is roman-numbered and the
    body restarts at arabic 1 at first_body_node_id — the same boundary the
    HTML/LaTeX surfaces switch on, so the three agree on where front matter ends.

    Args:
        data: the render-ready data dict.
        tree: optional per-session DocNode forest; None uses DOCUMENT_TREE.

    Returns:
        The .docx file contents as bytes (ready to write to disk or stream).
    """
    nodes = tree if tree is not None else DOCUMENT_TREE
    running_header = data.get("running_header") or data.get("title") or "5dToxReport"
    # Optional document-level style block (page geometry + base fonts); absent ⇒
    # the measured reference constants.  Same styles config the per-node layer
    # reads, under its own "document" key.
    document = (data.get("layout_style") or {}).get("document") or {}

    # Open the NTP styles-only base (386-style reference library, empty body) so
    # every style comes from the reference VERBATIM.  When the asset is absent,
    # fall back to a blank doc + the programmatic style skeleton (legacy path).
    doc, used_base = _open_base_document()
    if used_base:
        # The base supplies the reference fonts; honor an explicit styles.document
        # font override on top (geometry override is applied in _configure_section).
        _apply_document_font_overrides(doc, document)
    else:
        _build_style_skeleton(doc, document)

    # Role-driven styling (ADR-0010): resolve each role → NTP named style so the
    # content path applies real pStyles (3-02a_Head1_NoNumber, 0-03_Paragraph, …).
    # On the BASE path those styles already EXIST in the doc, so we DO NOT rebuild
    # them from the vocab graph (that would overwrite the reference's authentic
    # definitions with delta-derived approximations) — the vocab is used ONLY for
    # role→name resolution.  Default to the NTP vocab when the base is used; a
    # data-supplied `vocabulary` still wins.  On the legacy path, keep the prior
    # opt-in behaviour (build styles from the vocab only when explicitly named).
    if used_base and not data.get("vocabulary"):
        data = {**data, "vocabulary": _DEFAULT_DOCX_VOCAB}
    vocab = _load_active_vocabulary(data)
    if vocab is not None:
        if not used_base:
            # Legacy path: the blank doc has no NTP styles, so emit them.
            _build_vocabulary_styles(doc, vocab)
        # Stash on a shallow copy so handlers resolve a role → Word style name
        # (via _role_style_name) without threading a new parameter.
        data = {**data, "_vocabulary": vocab}
    front_section = doc.sections[0]
    _configure_section(front_section, running_header, document)
    # No page number on the cover page itself (front matter numbering still
    # counts it as page i, matching the reference's unnumbered cover).
    front_section.different_first_page_header_footer = True

    body_first_id = first_body_node_id(nodes)
    body_section_idx: int | None = None
    # FLOW state threaded across the WHOLE body walk (not reset per top-level
    # node): which orientation the current section is, and whether we've entered
    # the body region (front matter has no orientable nodes, so islands only
    # open in the body).  The body-start REGION section is portrait, matching the
    # initial landscape=False, so the first orientable node cleanly triggers the
    # first island.
    flow = {"landscape": False, "in_body": False}

    for top in nodes:
        if body_first_id is not None and top.id == body_first_id and not flow["in_body"]:
            # Section break → the body restarts page numbering at arabic 1.
            doc.add_section(WD_SECTION.NEW_PAGE)
            _configure_section(doc.sections[-1], running_header, document)
            # add_section() copies the front section's sectPr, including its
            # <w:titlePg> (different-first-page) flag — which the front section
            # sets ONLY to suppress the header on the cover.  The body has no
            # cover, so a distinct first page here just leaves an EMPTY, editable
            # first-page header on page 1 of the body (the stray "add header"
            # control the reviewer saw).  Turn it off so the body's first page
            # uses the same running header as the rest of the body.
            doc.sections[-1].different_first_page_header_footer = False
            # Capture the body-start section's POSITION.  FLOW (landscape)
            # sections open only inside the body and are always appended AFTER
            # this one, so its index stays valid — but a captured Section OBJECT
            # does not (add_section relocates sectPr elements between paragraphs,
            # so an object reference goes stale / points at the wrong region).
            # Index by final position at the end instead.
            body_section_idx = len(doc.sections) - 1
            flow["in_body"] = True
        _walk_docx_tree(doc, top, data, flow, document)

    # Page-number formats: front matter lower-roman, body decimal-from-1.
    # add_section() relocates the FIRST section's <w:sectPr> into the last
    # body paragraph, so a reference captured before that call goes stale —
    # re-fetch both sections from doc.sections by POSITION here, after all breaks
    # (REGION + FLOW islands) are in.  The body-start is NOT doc.sections[-1] once
    # landscape islands interleave — the last section is then a portrait resume —
    # so pin decimal to the captured body-start index, not the tail.
    _set_section_page_numbering(doc.sections[0], "lowerRoman", 1)
    if body_section_idx is not None:
        _set_section_page_numbering(doc.sections[body_section_idx], "decimal", 1)

    # A TOC (or any) field was emitted → tell the reader to recompute fields on
    # open so the contents fill in without a manual refresh.  Detect from the
    # rendered body rather than a data flag, so any future field-emitting handler
    # is covered without threading state through the walk.
    if doc.element.body.findall(".//" + qn("w:fldChar")):
        _mark_fields_dirty(doc)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
